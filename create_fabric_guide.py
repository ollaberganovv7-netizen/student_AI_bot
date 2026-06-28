from __future__ import annotations
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title ──
title = doc.add_heading('Crushed Velvet Fabric Texture — Substance Designer Node Guide', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════════════════════════════════════════
# STEP 1
# ═══════════════════════════════════════════════════════════
doc.add_heading('Step 1: Base Fiber Pattern (Horizontal Threads)', level=1)

doc.add_heading('Diagram:', level=3)
p = doc.add_paragraph()
run = p.add_run('[Anisotropic Noise] ──→ [Directional Warp] ──→ (Base Fibers)\n'
                '                              ↑\n'
                '                     [Perlin Noise]')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Node Settings:', level=3)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Node'
hdr[1].text = 'Parameter'
hdr[2].text = 'Value'
for row_data in [
    ('Anisotropic Noise', 'Direction', '0° (horizontal)'),
    ('', 'Stretching', '0.85'),
    ('', 'Amount', '12'),
    ('', 'Output size', '2048×2048'),
    ('Perlin Noise', 'Scale', '5'),
    ('', 'Disorder', '0.3'),
    ('Directional Warp', 'Input', 'Anisotropic output'),
    ('', 'Intensity input', 'Perlin output'),
    ('', 'Warp Angle', '0°'),
    ('', 'Intensity', '3.0'),
]:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# ═══════════════════════════════════════════════════════════
# STEP 2
# ═══════════════════════════════════════════════════════════
doc.add_heading('Step 2: Crushed Velvet Variation (Mottled Look)', level=1)

doc.add_heading('Diagram:', level=3)
p = doc.add_paragraph()
run = p.add_run('[Perlin Noise Zoom] ──→ [Levels 1] ──→ [Blend: Multiply] ──→ (Crush Map)\n'
                '                                              ↑\n'
                '[Clouds 2] ──────────→ [Levels 2] ───────────┘')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Node Settings:', level=3)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Node'
hdr[1].text = 'Parameter'
hdr[2].text = 'Value'
for row_data in [
    ('Perlin Noise Zoom', 'Scale', '3'),
    ('', 'Disorder', '0.5'),
    ('', 'Brightness', '0.6'),
    ('Levels 1', 'Level In Low', '0.25'),
    ('', 'Level In High', '0.85'),
    ('Clouds 2', 'Scale', '6'),
    ('', 'Disorder', '1.0'),
    ('Levels 2', 'Level In Low', '0.3'),
    ('', 'Level In High', '0.7'),
    ('Blend', 'Foreground', 'Levels 1 output'),
    ('', 'Background', 'Levels 2 output'),
    ('', 'Blending Mode', 'Multiply'),
    ('', 'Opacity', '0.7'),
]:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# ═══════════════════════════════════════════════════════════
# STEP 3
# ═══════════════════════════════════════════════════════════
doc.add_heading('Step 3: Combine Fibers + Crush', level=1)

doc.add_heading('Diagram:', level=3)
p = doc.add_paragraph()
run = p.add_run('(Base Fibers) ──→ [Blend: Multiply] ──→ [Levels] ──→ (Height Map)\n'
                '                        ↑\n'
                '              (Crush Map) ──────────┘')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Node Settings:', level=3)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Node'
hdr[1].text = 'Parameter'
hdr[2].text = 'Value'
for row_data in [
    ('Blend', 'Foreground', 'Base Fibers'),
    ('', 'Background', 'Crush Map'),
    ('', 'Blending Mode', 'Multiply'),
    ('', 'Opacity', '0.6'),
    ('Levels', 'Level In Low', '0.15'),
    ('', 'Level In High', '0.9'),
    ('', 'Level Out Low', '0.1'),
    ('', 'Level Out High', '0.85'),
]:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# ═══════════════════════════════════════════════════════════
# STEP 4
# ═══════════════════════════════════════════════════════════
doc.add_heading('Step 4: Normal Map', level=1)

doc.add_heading('Diagram:', level=3)
p = doc.add_paragraph()
run = p.add_run('(Height Map) ──→ [Normal] ──→ Output: Normal')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Node Settings:', level=3)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Node'
hdr[1].text = 'Parameter'
hdr[2].text = 'Value'
for row_data in [
    ('Normal', 'Input', 'Height Map'),
    ('', 'Intensity', '4.0'),
    ('', 'Format', 'DirectX or OpenGL (your engine)'),
]:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# ═══════════════════════════════════════════════════════════
# STEP 5
# ═══════════════════════════════════════════════════════════
doc.add_heading('Step 5: Base Color (Dark Olive Green)', level=1)

doc.add_heading('Diagram:', level=3)
p = doc.add_paragraph()
run = p.add_run('(Height Map) ──→ [Gradient Map] ──→ [HSL] ──→ Output: Base Color')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Node Settings:', level=3)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Node'
hdr[1].text = 'Parameter'
hdr[2].text = 'Value'
for row_data in [
    ('Gradient Map', 'Input', 'Height Map'),
    ('', 'Left color (dark)', '#1A1F0E (very dark olive)'),
    ('', 'Mid color', '#2E3318 (dark green)'),
    ('', 'Right color (light)', '#4A4F28 (olive highlight)'),
    ('', 'Mid position', '0.45'),
    ('HSL', 'Saturation', '-0.05'),
    ('', 'Lightness', '-0.1'),
]:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# ═══════════════════════════════════════════════════════════
# STEP 6
# ═══════════════════════════════════════════════════════════
doc.add_heading('Step 6: Roughness Map', level=1)

doc.add_heading('Diagram:', level=3)
p = doc.add_paragraph()
run = p.add_run('(Height Map) ──→ [Invert] ──→ [Levels] ──→ Output: Roughness')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Node Settings:', level=3)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Node'
hdr[1].text = 'Parameter'
hdr[2].text = 'Value'
for row_data in [
    ('Invert', 'Input', 'Height Map'),
    ('Levels', 'Level Out Low', '0.55'),
    ('', 'Level Out High', '0.85'),
]:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

p = doc.add_paragraph()
run = p.add_run('Note: Velvet has HIGH roughness overall; crushed areas are slightly shinier.')
run.italic = True

# ═══════════════════════════════════════════════════════════
# STEP 7
# ═══════════════════════════════════════════════════════════
doc.add_heading('Step 7: AO (Ambient Occlusion)', level=1)

doc.add_heading('Diagram:', level=3)
p = doc.add_paragraph()
run = p.add_run('(Height Map) ──→ [Ambient Occlusion (HBAO)] ──→ Output: AO')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Node Settings:', level=3)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Node'
hdr[1].text = 'Parameter'
hdr[2].text = 'Value'
for row_data in [
    ('HBAO', 'Height input', 'Height Map'),
    ('', 'Radius', '0.75'),
    ('', 'Quality', '8'),
]:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# ═══════════════════════════════════════════════════════════
# FULL GRAPH
# ═══════════════════════════════════════════════════════════
doc.add_heading('Full Graph Overview', level=1)

p = doc.add_paragraph()
run = p.add_run(
    '[Anisotropic Noise] → [Dir. Warp] ──────────────┐\n'
    '        ↑                                        ↓\n'
    '[Perlin Noise]                             [Blend: Multiply] → [Levels] → HEIGHT\n'
    '                                                 ↑                           │\n'
    '[Perlin Noise Zoom] → [Levels] ──┐              │                           ├→ [Normal] → NORMAL\n'
    '                                  ↓              │                           ├→ [Gradient Map] → [HSL] → BASE COLOR\n'
    '                            [Blend: Multiply] ───┘                           ├→ [Invert] → [Levels] → ROUGHNESS\n'
    '                                  ↑                                          └→ [HBAO] → AO\n'
    '                   [Clouds 2] → [Levels] ──┘'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════
# OUTPUT SUMMARY
# ═══════════════════════════════════════════════════════════
doc.add_heading('Output Summary', level=1)

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Output'
hdr[1].text = 'Source'
for row_data in [
    ('Base Color', 'Gradient Map → HSL'),
    ('Normal', 'Normal node (intensity 4.0)'),
    ('Height', 'Blended fibers × crush'),
    ('Roughness', 'Inverted height, range 0.55–0.85'),
    ('AO', 'HBAO from height'),
]:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

p = doc.add_paragraph()
run = p.add_run(
    'KEY: The main effect comes from Anisotropic Noise (horizontal fibers) '
    'multiplied with Perlin + Clouds (crushed mottling). Adjust the Gradient '
    'Map colors to dial in the exact olive tone.'
)
run.italic = True

doc.save('/Users/arslon/Desktop/student_bot/Fabric_Texture_Guide.docx')
print("Saved: Fabric_Texture_Guide.docx")
