from __future__ import annotations
"""
Kurs ishi (course work) DOCX generator.
Structure based on: baxtiyor.uz/kurs-ishi-yozish-...

Generates a properly formatted DOCX with:
- Titul varog'i
- Reja (with dotted page numbers)
- Kirish
- Asosiy qism (2 bob, 3 paragraf each)
- Xulosa va takliflar
- Foydalanilgan adabiyotlar
- Mundarija

Usage:
    from services.kurs_ishi_generator import generate_kurs_ishi
    buf = generate_kurs_ishi(
        mavzu="...",
        fan="...",
        talaba="...",
        guruh="...",
        rahbar="...",
        bob1_nomi="...",
        bob1_paragraflar=["1.1...", "1.2...", "1.3..."],
        bob2_nomi="...",
        bob2_paragraflar=["2.1...", "2.2...", "2.3..."],
    )
"""
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _add_paragraph(doc, text, font_size=14, bold=False, italic=False,
                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(0),
                   space_before=Pt(0), first_line_indent=Cm(1.25),
                   font_name="Times New Roman"):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = 1.5
    if first_line_indent:
        pf.first_line_indent = first_line_indent

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Set font for Cyrillic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)
    return p


def _add_empty_paragraph(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.5
        run = p.add_run("")
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"


def _add_page_break(doc):
    doc.add_page_break()


def _set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def _make_reja_line(name, page_num, total_width=60):
    """Create a line like: 1.1. Name ............... 5"""
    dots_needed = total_width - len(name) - len(str(page_num)) - 2
    if dots_needed < 3:
        dots_needed = 3
    return f"{name} {'.' * dots_needed} {page_num}"


def _add_titul(doc, mavzu, fan, talaba, guruh, rahbar, universitet=""):
    """Add title page."""
    if not universitet:
        universitet = "O'ZBEKISTON RESPUBLIKASI\nRaqamli texnologiyalar vazirligi\n\nMUHAMMAD AL-XORAZMIY NOMIDAGI\nTOSHKENT AXBOROT TEXNOLOGIYALARI UNIVERSITETI"

    # University name
    for line in universitet.split("\n"):
        _add_paragraph(
            doc, line,
            font_size=14, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=None,
            space_after=Pt(2),
        )

    _add_empty_paragraph(doc, 3)

    # Fan nomi
    _add_paragraph(
        doc, f'"{fan}" fanidan',
        font_size=14, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
    )

    _add_empty_paragraph(doc, 1)

    # KURS ISHI
    _add_paragraph(
        doc, "KURS ISHI",
        font_size=18, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
    )

    _add_empty_paragraph(doc, 1)

    # Mavzu
    _add_paragraph(
        doc, f"Mavzu: {mavzu}",
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
    )

    _add_empty_paragraph(doc, 4)

    # Talaba va rahbar info (right-aligned)
    _add_paragraph(
        doc, f"Bajardi: {talaba}",
        font_size=14, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=None,
    )
    _add_paragraph(
        doc, f"Guruh: {guruh}",
        font_size=14, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=None,
    )
    _add_paragraph(
        doc, f"Ilmiy rahbar: {rahbar}",
        font_size=14, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=None,
    )

    _add_empty_paragraph(doc, 3)

    # City and year
    _add_paragraph(
        doc, "Toshkent — 2025",
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
    )


def _add_reja(doc, bob1_nomi, bob1_paragraflar, bob2_nomi, bob2_paragraflar):
    """Add REJA page with dotted page numbers."""
    _add_paragraph(
        doc, "REJA:",
        font_size=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=Pt(12),
    )

    _add_empty_paragraph(doc, 1)

    # KIRISH
    _add_paragraph(
        doc, _make_reja_line("KIRISH", 3),
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        space_after=Pt(6),
    )

    _add_empty_paragraph(doc, 1)

    # Bob 1
    _add_paragraph(
        doc, f"1.Bob. {bob1_nomi}",
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Cm(1),
        space_after=Pt(4),
    )

    pages_1 = [6, 9, 12]
    for i, par in enumerate(bob1_paragraflar[:3]):
        _add_paragraph(
            doc, _make_reja_line(f"    1.{i+1}. {par}", pages_1[i]),
            font_size=14, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Cm(1.5),
            space_after=Pt(2),
        )

    _add_empty_paragraph(doc, 1)

    # Bob 2
    _add_paragraph(
        doc, f"2.Bob. {bob2_nomi}",
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Cm(1),
        space_after=Pt(4),
    )

    pages_2 = [15, 18, 21]
    for i, par in enumerate(bob2_paragraflar[:3]):
        _add_paragraph(
            doc, _make_reja_line(f"    2.{i+1}. {par}", pages_2[i]),
            font_size=14, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Cm(1.5),
            space_after=Pt(2),
        )

    _add_empty_paragraph(doc, 1)

    # Xulosa
    _add_paragraph(
        doc, _make_reja_line("XULOSA VA TAKLIFLAR", 24),
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        space_after=Pt(4),
    )

    # Adabiyotlar
    _add_paragraph(
        doc, _make_reja_line("FOYDALANILGAN ADABIYOTLAR", 27),
        font_size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        space_after=Pt(4),
    )


def generate_kurs_ishi(
    mavzu: str,
    fan: str = "Informatika",
    talaba: str = "Talaba F.I.Sh.",
    guruh: str = "000-00",
    rahbar: str = "Rahbar F.I.Sh.",
    universitet: str = "",
    bob1_nomi: str = "",
    bob1_paragraflar: list = None,
    bob2_nomi: str = "",
    bob2_paragraflar: list = None,
) -> io.BytesIO:
    """Generate a kurs ishi DOCX and return as BytesIO."""

    if not bob1_paragraflar:
        bob1_paragraflar = [
            "Mavzuning nazariy asoslari",
            "Asosiy tushunchalar va ta'riflar",
            "Mavzuning dolzarbligi va ahamiyati",
        ]
    if not bob2_paragraflar:
        bob2_paragraflar = [
            "Amaliy tahlil va natijalar",
            "Mavjud muammolar va yechimlar",
            "Taklif etiladigan yondashuvlar",
        ]
    if not bob1_nomi:
        bob1_nomi = f"{mavzu}: nazariy asoslar"
    if not bob2_nomi:
        bob2_nomi = f"{mavzu}: amaliy tahlil"

    doc = Document()
    _set_margins(doc)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    # ────── TITUL VAROG'I ──────
    _add_titul(doc, mavzu, fan, talaba, guruh, rahbar, universitet)

    _add_page_break(doc)

    # ────── REJA ──────
    _add_reja(doc, bob1_nomi, bob1_paragraflar, bob2_nomi, bob2_paragraflar)

    _add_page_break(doc)

    # ────── KIRISH ──────
    _add_paragraph(
        doc, "KIRISH",
        font_size=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=Pt(12),
    )

    kirish_text = (
        f"Bugungi kunda {mavzu.lower()} masalasi tobora dolzarb bo'lib bormoqda. "
        f"Zamonaviy dunyoda ushbu sohaning rivojlanishi katta ahamiyatga ega. "
        f"Mazkur kurs ishida {mavzu.lower()} mavzusi atroflicha o'rganiladi.\n\n"
        f"Kurs ishining maqsadi — {mavzu.lower()} bo'yicha nazariy va amaliy "
        f"bilimlarni chuqurlashtirish, mavjud muammolarni aniqlash va ularning "
        f"yechimlarini taklif etishdan iborat.\n\n"
        f"Kurs ishining vazifalari:\n"
        f"— mavzuning nazariy asoslarini o'rganish;\n"
        f"— mavjud adabiyotlar va manbalarni tahlil qilish;\n"
        f"— amaliy ma'lumotlarni yig'ish va qayta ishlash;\n"
        f"— olingan natijalarni umumlashtirish;\n"
        f"— xulosa va takliflar ishlab chiqish.\n\n"
        f"Kurs ishining ob'ekti — {mavzu.lower()} sohasidagi jarayonlar va hodisalar. "
        f"Kurs ishining predmeti esa ushbu sohadagi nazariy va amaliy jihatlardir.\n\n"
        f"Kurs ishi kirish, ikki bob, xulosa va foydalanilgan adabiyotlar ro'yxatidan "
        f"tashkil topgan."
    )

    for para in kirish_text.split("\n\n"):
        if para.strip():
            _add_paragraph(doc, para.strip(), font_size=14, space_after=Pt(6))

    _add_page_break(doc)

    # ────── BOB 1 ──────
    _add_paragraph(
        doc, f"1-BOB. {bob1_nomi.upper()}",
        font_size=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=Pt(12),
    )

    for i, par_name in enumerate(bob1_paragraflar[:3]):
        _add_paragraph(
            doc, f"1.{i+1}. {par_name}",
            font_size=14, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=None,
            space_after=Pt(8),
            space_before=Pt(12),
        )

        # Placeholder text for each paragraph
        placeholder = (
            f"Ushbu bo'limda {par_name.lower()} masalasi ko'rib chiqiladi. "
            f"[Bu yerga {par_name.lower()} bo'yicha batafsil ma'lumot yoziladi. "
            f"Hajmi kamida 2-3 sahifa bo'lishi kerak.]\n\n"
            f"Mavzuning ushbu jihati haqida turli mualliflar o'z fikrlarini "
            f"bildirganlar. Masalan, ... [manbalardan iqtiboslar keltiring].\n\n"
            f"Shunday qilib, {par_name.lower()} masalasi bo'yicha quyidagi "
            f"xulosaga kelish mumkin: ..."
        )

        for para in placeholder.split("\n\n"):
            if para.strip():
                _add_paragraph(doc, para.strip(), font_size=14, space_after=Pt(6))

        if i < 2:
            _add_empty_paragraph(doc, 1)

    _add_page_break(doc)

    # ────── BOB 2 ──────
    _add_paragraph(
        doc, f"2-BOB. {bob2_nomi.upper()}",
        font_size=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=Pt(12),
    )

    for i, par_name in enumerate(bob2_paragraflar[:3]):
        _add_paragraph(
            doc, f"2.{i+1}. {par_name}",
            font_size=14, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=None,
            space_after=Pt(8),
            space_before=Pt(12),
        )

        placeholder = (
            f"Ushbu bo'limda {par_name.lower()} masalasi amaliy jihatdan "
            f"tahlil qilinadi. [Bu yerga {par_name.lower()} bo'yicha batafsil "
            f"ma'lumot, jadvallar, grafiklar va tahlil natijalari yoziladi. "
            f"Hajmi kamida 2-3 sahifa bo'lishi kerak.]\n\n"
            f"Olingan ma'lumotlar tahlili shuni ko'rsatadiki, ... "
            f"[statistik ma'lumotlar va dalillar keltiring].\n\n"
            f"Yuqoridagilardan kelib chiqib, {par_name.lower()} bo'yicha "
            f"quyidagi natijalarni ta'kidlash mumkin: ..."
        )

        for para in placeholder.split("\n\n"):
            if para.strip():
                _add_paragraph(doc, para.strip(), font_size=14, space_after=Pt(6))

        if i < 2:
            _add_empty_paragraph(doc, 1)

    _add_page_break(doc)

    # ────── XULOSA VA TAKLIFLAR ──────
    _add_paragraph(
        doc, "XULOSA VA TAKLIFLAR",
        font_size=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=Pt(12),
    )

    xulosa_text = (
        f"Mazkur kurs ishida {mavzu.lower()} mavzusi bo'yicha nazariy va amaliy "
        f"tadqiqot olib borildi. Tadqiqot natijasida quyidagi xulosalarga kelindi:\n\n"
        f"Birinchidan, {bob1_paragraflar[0].lower()} bo'yicha ... "
        f"[1-bob natijalari asosida xulosa].\n\n"
        f"Ikkinchidan, {bob2_paragraflar[0].lower()} bo'yicha ... "
        f"[2-bob natijalari asosida xulosa].\n\n"
        f"Takliflar:\n"
        f"1. [Birinchi taklif]\n"
        f"2. [Ikkinchi taklif]\n"
        f"3. [Uchinchi taklif]\n\n"
        f"Umid qilamizki, ushbu tadqiqot natijalari {mavzu.lower()} sohasida "
        f"amaliy qo'llaniladi va keyingi ilmiy izlanishlar uchun asos bo'ladi."
    )

    for para in xulosa_text.split("\n\n"):
        if para.strip():
            _add_paragraph(doc, para.strip(), font_size=14, space_after=Pt(6))

    _add_page_break(doc)

    # ────── FOYDALANILGAN ADABIYOTLAR ──────
    _add_paragraph(
        doc, "FOYDALANILGAN ADABIYOTLAR",
        font_size=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=Pt(12),
    )

    adabiyotlar = [
        "O'zbekiston Respublikasi Prezidentining farmonlari va qarorlari.",
        f"{mavzu} bo'yicha o'quv qo'llanma. — T.: O'qituvchi, 2023.",
        "Karimov I.A. Yuksak ma'naviyat — yengilmas kuch. — T.: Ma'naviyat, 2008.",
        f"{mavzu}: nazariya va amaliyot. — T.: Fan, 2022.",
        "Xolmatov T.X. Informatika asoslari. — T.: O'qituvchi, 2021.",
        "Internet manba: www.ziyonet.uz",
        "Internet manba: www.scholar.google.com",
        "Internet manba: www.edu.uz",
    ]

    for i, adb in enumerate(adabiyotlar):
        _add_paragraph(
            doc, f"{i+1}. {adb}",
            font_size=14, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=None,
            space_after=Pt(4),
        )

    # Save to BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
