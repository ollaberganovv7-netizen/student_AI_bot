from __future__ import annotations
import io, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_LEADER, WD_TAB_ALIGNMENT

SERVICE_TITLES = {
    "essay": "REFERAT",
    "report": "MUSTAQIL ISH",
    "coursework": "KURS ISHI",
    "tezis": "ILMIY MAQOLA (TEZIS)",
    "maqola": "ILMIY MAQOLA",
}

def _set_heading_style(paragraph, level: int = 1):
    """Apply standard academic heading formatting."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(16)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        run.font.size = Pt(14)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(12)

def _reja_line(name: str, page: int, width: int = 55) -> str:
    """Format: Name ............... page_num"""
    dots = width - len(name) - len(str(page)) - 2
    if dots < 3:
        dots = 3
    return f"{name} {'.' * dots} {page}"


def _add_reja_item(doc, text: str, bold: bool = False, indent: float = 0):
    """Add a single REJA line with optional bold and indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _build_reja_page(doc, plan: str):
    """Build a properly formatted REJA page with dotted page numbers."""
    # Title
    p_title = doc.add_paragraph("REJA:")
    _set_heading_style(p_title, level=1)
    p_title.paragraph_format.first_line_indent = Cm(0)
    p_title.paragraph_format.space_after = Pt(12)

    doc.add_paragraph("")  # spacer

    # Parse plan into sections
    plan_lines = [s.strip() for s in plan.split("\n") if s.strip()]

    # Classify lines: KIRISH, bob headers, sub-sections, XULOSA, ADABIYOTLAR
    page_counter = 3  # start from page 3

    for line in plan_lines:
        line_upper = line.upper().strip()
        # Strip leading numbering
        clean = re.sub(r'^[\d\.]+\s*', '', line).strip()

        if "KIRISH" in line_upper:
            _add_reja_item(doc, _reja_line("KIRISH", page_counter), bold=True)
            page_counter += 2

        elif "XULOSA" in line_upper:
            doc.add_paragraph("")
            _add_reja_item(doc, _reja_line("XULOSA VA TAKLIFLAR", page_counter), bold=True)
            page_counter += 2

        elif "ADABIYOTLAR" in line_upper or "FOYDALANILGAN" in line_upper:
            _add_reja_item(doc, _reja_line("FOYDALANILGAN ADABIYOTLAR", page_counter), bold=True)

        elif "ASOSIY QISM" in line_upper or re.match(r'^\d+[\.\-]\s*[Bb]ob', line) or re.match(r'^[IVX]+[\.\-]?\s*[Bb]ob', line):
            # Bob header (e.g., "1-Bob. Nazariy asoslar" or "I Bob. ...")
            doc.add_paragraph("")
            _add_reja_item(doc, line.strip(), bold=True, indent=0.5)

        elif re.match(r'^\d+\.\d+', line):
            # Sub-section (e.g., "1.1. Topic name")
            _add_reja_item(doc, _reja_line(f"    {line.strip()}", page_counter), indent=1.0)
            page_counter += 3

        else:
            # Generic section
            _add_reja_item(doc, _reja_line(clean, page_counter), indent=0.5)
            page_counter += 3


def _build_mundarija_table(doc, plan: str, total_pages: int = 15):
    """Build a properly formatted MUNDARIJA page with a bordered table,
    dotted leaders and page numbers — matching the academic standard.
    Page numbers are distributed proportionally based on total_pages."""
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    # Title
    p_title = doc.add_paragraph("MUNDARIJA")
    _set_heading_style(p_title, level=1)
    p_title.paragraph_format.first_line_indent = Cm(0)
    p_title.paragraph_format.space_after = Pt(6)

    # Parse plan into structured items
    plan_lines = [s.strip() for s in plan.split("\n") if s.strip()]
    raw_items = []  # list of (text, is_bold, is_bob_header, weight)

    for line in plan_lines:
        line_upper = line.upper().strip()

        if "KIRISH" in line_upper:
            raw_items.append(("KIRISH", True, False, 1.0))

        elif "XULOSA" in line_upper:
            raw_items.append(("XULOSA", True, False, 1.0))

        elif "ADABIYOTLAR" in line_upper or "FOYDALANILGAN" in line_upper:
            raw_items.append(("FOYDALANILGAN ADABIYOTLAR", True, False, 0.5))

        elif re.match(r'^[IVX]+[\.\-]?\s*[Bb][Oo][Bb]', line) or re.match(r'^\d+[\.\-]\s*[Bb]ob', line):
            raw_items.append((line.strip(), True, True, 0))  # bob header shares page with first subsection

        elif re.match(r'^\d+\.\d+', line):
            raw_items.append((line.strip(), False, False, 1.0))

        else:
            clean = re.sub(r'^[\d\.]+\s*', '', line).strip()
            if clean:
                raw_items.append((clean, False, False, 1.0))

    if not raw_items:
        return

    # Calculate page numbers proportionally
    # Pages: cover(1) + mundarija(1) + content(total_pages - 2)
    content_pages = max(total_pages - 2, 3)  # at least 3 content pages
    
    # Sum all weights to distribute content pages
    total_weight = sum(w for _, _, _, w in raw_items)
    if total_weight <= 0:
        total_weight = 1

    # Assign pages proportionally
    items = []
    current_page = 3  # content starts at page 3 (after cover + mundarija)
    
    for text, is_bold, is_bob, weight in raw_items:
        page = int(current_page)
        items.append((text, page, is_bold, is_bob))
        # Advance page counter proportionally
        if weight > 0:
            pages_for_section = (weight / total_weight) * content_pages
            current_page += max(pages_for_section, 1.0)

    # Create table with borders
    table = doc.add_table(rows=len(items), cols=2)
    table.style = 'Table Grid'

    # Set column widths
    for row_idx, (text, page, is_bold, is_bob) in enumerate(items):
        row = table.rows[row_idx]
        # Name cell
        cell_name = row.cells[0]
        cell_name.width = Cm(14.5)
        p = cell_name.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        if is_bob:
            # BOB headers: bold, slightly larger
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.bold = True
        elif is_bold:
            # KIRISH, XULOSA, ADABIYOTLAR
            dots_count = 60 - len(text)
            if dots_count < 3:
                dots_count = 3
            run = p.add_run(f"{text}{'.' * dots_count}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.bold = True
        else:
            # Sub-sections with dotted leaders
            dots_count = 55 - len(text)
            if dots_count < 3:
                dots_count = 3
            run = p.add_run(f"{text}{'.' * dots_count}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        # Page number cell
        cell_page = row.cells[1]
        cell_page.width = Cm(1.5)
        p_page = cell_page.paragraphs[0]
        p_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_page.paragraph_format.first_line_indent = Cm(0)
        p_page.paragraph_format.space_before = Pt(2)
        p_page.paragraph_format.space_after = Pt(2)
        run_page = p_page.add_run(str(page))
        run_page.font.name = "Times New Roman"
        run_page.font.size = Pt(12)
        if is_bold:
            run_page.font.bold = True

def generate_docx(service_type: str, topic: str, content: str, author: str = "Talaba", plan: str = "", extra_meta: dict = None) -> bytes:
    doc = Document()
    extra_meta = extra_meta or {}

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)

    # ── Default style ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.first_line_indent = Cm(1.25)

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    if service_type == "coursework":
        # Professional Coursework Cover
        min_p = doc.add_paragraph(extra_meta.get("ministry", "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM VAZIRLIGI").upper())
        min_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        min_p.paragraph_format.first_line_indent = Cm(0)
        min_p.runs[0].font.size = Pt(12)
        min_p.runs[0].font.bold = True

        for _ in range(3): doc.add_paragraph("")

        kaf_p = doc.add_paragraph(extra_meta.get("department", "KAFEDRA NOMI").upper())
        kaf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kaf_p.paragraph_format.first_line_indent = Cm(0)
        kaf_p.runs[0].font.size = Pt(14)
        kaf_p.runs[0].font.bold = True

        for _ in range(4): doc.add_paragraph("")

        p_type = doc.add_paragraph("KURS ISHI")
        p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_type.paragraph_format.first_line_indent = Cm(0)
        p_type.runs[0].font.size = Pt(28)
        p_type.runs[0].font.bold = True

        p_top = doc.add_paragraph(f"MAVZU: {topic.upper()}")
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_top.paragraph_format.first_line_indent = Cm(0)
        p_top.runs[0].font.size = Pt(16)
        p_top.runs[0].font.italic = True

        for _ in range(6): doc.add_paragraph("")

        p_auth = doc.add_paragraph()
        p_auth.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_auth.paragraph_format.first_line_indent = Cm(0)
        run_auth = p_auth.add_run(f"Bajardi: {author}\nTekshirdi: ________________")
        run_auth.font.size = Pt(14)
        run_auth.font.bold = True

    elif service_type == "tezis":
        # Tezis title page: topic bold caps centered, author right italic
        p_top = doc.add_paragraph(topic.upper())
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_top.paragraph_format.first_line_indent = Cm(0)
        p_top.runs[0].font.size = Pt(16)
        p_top.runs[0].font.bold = True

        p_auth = doc.add_paragraph(f"Muallif: {author}")
        p_auth.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_auth.paragraph_format.first_line_indent = Cm(0)
        p_auth.runs[0].font.size = Pt(12)
        p_auth.runs[0].font.italic = True

        doc.add_paragraph("")  # Space before text

    elif service_type == "maqola":
        # Maqola title page: matches academic journal format
        # 1) Title — bold, caps, centered
        p_title = doc.add_paragraph(topic.upper())
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.first_line_indent = Cm(0)
        p_title.paragraph_format.space_after = Pt(6)
        run_title = p_title.runs[0]
        run_title.font.name = "Times New Roman"
        run_title.font.size = Pt(14)
        run_title.font.bold = True

        # 2) Author — bold, centered
        p_auth = doc.add_paragraph(author)
        p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_auth.paragraph_format.first_line_indent = Cm(0)
        p_auth.paragraph_format.space_after = Pt(2)
        run_auth = p_auth.runs[0]
        run_auth.font.name = "Times New Roman"
        run_auth.font.size = Pt(14)
        run_auth.font.bold = True

        # 3) University — italic, centered (only if provided)
        university = extra_meta.get("university", "")
        if university:
            p_uni = doc.add_paragraph(university)
            p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_uni.paragraph_format.first_line_indent = Cm(0)
            p_uni.paragraph_format.space_after = Pt(2)
            run_uni = p_uni.runs[0]
            run_uni.font.name = "Times New Roman"
            run_uni.font.size = Pt(14)
            run_uni.font.italic = True

        doc.add_paragraph("")  # Space before content

    elif service_type == "uslubiy":
        # Load from uslubiy template if it exists
        import os
        template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "doc_templates", "uslubiy_template.docx")
        if os.path.exists(template_path):
            doc = Document(template_path)
            university = extra_meta.get("university", "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM VAZIRLIGI").upper()
            import datetime
            year = str(datetime.datetime.now().year)

            for p in doc.paragraphs:
                text = p.text
                if "Universitet" in text and "___" in text:
                    for run in p.runs:
                        if "___" in run.text:
                            run.text = run.text.replace("_______________________________________________", university)
                            break
                elif "FAN" in text and "___" in text:
                    for run in p.runs:
                        if "___" in run.text:
                            subject = extra_meta.get("subject", "")
                            if subject:
                                run.text = run.text.replace("_________________", subject.upper())
                            break
                elif "1-MAVZU:" in text and "___" in text:
                    for run in p.runs:
                        if "___" in run.text:
                            run.text = run.text.replace("____________________________________", topic.upper())
                            break
                elif "202" in text and "yil" in text:
                    for run in p.runs:
                        if "202" in run.text:
                            run.text = run.text.replace("2026", year)
                            break
        else:
            # Fallback if template is missing
            doc.add_paragraph("USLUBIY ISHLANMA").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"MAVZU: {topic.upper()}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    else:
        # Standard Cover (Referat/Mustaqil ish)
        doc_type_label = SERVICE_TITLES.get(service_type, "REFERAT")
        p1 = doc.add_paragraph("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM VAZIRLIGI")
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.first_line_indent = Cm(0)
        p1.runs[0].font.size = Pt(12)
        p1.runs[0].font.bold = True
        for _ in range(5): doc.add_paragraph("")
        p2 = doc.add_paragraph(doc_type_label)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.first_line_indent = Cm(0)
        p2.runs[0].font.size = Pt(28)
        p2.runs[0].font.bold = True
        p3 = doc.add_paragraph(f"Mavzu: {topic}")
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.first_line_indent = Cm(0)
        p3.runs[0].font.size = Pt(16)
        p3.runs[0].font.italic = True
        for _ in range(8): doc.add_paragraph("")
        p_auth = doc.add_paragraph()
        p_auth.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_auth.paragraph_format.first_line_indent = Cm(0)
        run_auth = p_auth.add_run(f"Bajardi: {formatted_author}")
        run_auth.font.size = Pt(14)
        run_auth.font.bold = True

    if service_type != "tezis":
        doc.add_page_break()

    # ── MUNDARIJA / REJA ─────────────────────────────────────────────────
    if plan:
        if service_type == "coursework":
            _build_reja_page(doc, plan)
        else:
            total_pgs = extra_meta.get("num_pages", 15) if extra_meta else 15
            _build_mundarija_table(doc, plan, total_pages=total_pgs)
        doc.add_page_break()

    # ── CONTENT ─────────────────────────────────────────────────────────────
    lines = content.split("\n")
    i_line = 0
    while i_line < len(lines):
        stripped = lines[i_line].strip()
        if not stripped:
            i_line += 1
            continue

        if stripped.startswith("#"):
            level = 1 if stripped.startswith("# ") else 2
            p = doc.add_paragraph(stripped.lstrip("#").strip())
            _set_heading_style(p, level=level)
            p.paragraph_format.first_line_indent = Cm(0)
            i_line += 1
            continue

        # ── Image placeholder: [ 🖼️ SHU YERGA RASM JOYLANG: ... ]
        if "🖼" in stripped and "SHU YERGA RASM JOYLANG" in stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(stripped)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            i_line += 1
            continue

        # ── Table: JADVAL: header followed by | delimited rows
        if stripped.upper().startswith("JADVAL:") or ("|" in stripped and stripped.count("|") >= 2):
            # Collect all table rows
            table_rows = []
            if stripped.upper().startswith("JADVAL:"):
                i_line += 1  # skip "JADVAL:" label
            while i_line < len(lines):
                row_text = lines[i_line].strip()
                if not row_text or ("|" not in row_text and row_text.count("|") < 1):
                    break
                # Skip separator lines like |---|---|
                if re.match(r'^[\s|:-]+$', row_text):
                    i_line += 1
                    continue
                cells = [c.strip() for c in row_text.split("|") if c.strip() != ""]
                if cells:
                    table_rows.append(cells)
                i_line += 1

            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                tbl.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < max_cols:
                            cell = tbl.cell(r_idx, c_idx)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.first_line_indent = Cm(0)
                                for run in paragraph.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(12)
                                    if r_idx == 0:
                                        run.bold = True
                doc.add_paragraph("")  # Space after table
            continue

        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        i_line += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Template-based DOCX generation (Referat) ────────────────────────────────

import os

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "doc_templates")

def generate_docx_from_template(
    topic: str,
    content: str,
    author: str = "Talaba",
    plan: str = "",
    subject: str = "",
    reviewer: str = "",
    university: str = "",
    fakultet: str = "",
    specialty: str = "",
    doc_type: str = "referat",
    num_pages: int = 15
) -> bytes:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
    from docx import Document
    import io, os, re
    from docx.shared import Pt
    
    
    def format_fio(full_name: str) -> str:
        parts = full_name.strip().split()
        if len(parts) == 0:
            return ""
        if len(parts) == 1:
            return parts[0].capitalize()
        
        last_name = parts[0].capitalize()
        first_init = parts[1][0].upper() + "." if len(parts) > 1 else ""
        middle_init = parts[2][0].upper() + "." if len(parts) > 2 else ""
        return f"{last_name} {first_init} {middle_init}".strip()
        
    formatted_author = format_fio(author)
    
    if doc_type in ["report", "mustaqil", "mustaqil_ish"]:
        template_path = os.path.join(TEMPLATE_DIR, "Mustaqil ish_yaratish.docx")
    else:
        template_path = os.path.join(TEMPLATE_DIR, "Referat_yaratish.docx")
        
    doc = Document(template_path)
    
    # Hide page number on the first page
    doc.sections[0].different_first_page_header_footer = True
    
    # Remove 2 blank paragraphs before "Tashkent-2026" to prevent it from moving to the next page
    blanks_deleted = 0
    found_qabul = False
    for p in list(doc.paragraphs[:30]):
        if "Qabul qildi" in p.text:
            found_qabul = True
            continue
        if found_qabul and not p.text.strip():
            p._element.getparent().remove(p._element)
            blanks_deleted += 1
            if blanks_deleted >= 2:
                break
        elif found_qabul and p.text.strip():
            break

    # 0. Apply official academic formatting margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)

    # 1. Parse AI Content
    sections_list = []
    current_section = None
    current_lines = []

    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            if current_section is not None:
                sections_list.append((current_section, "\n".join(current_lines)))
            current_section = stripped[2:].strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_section is not None:
        sections_list.append((current_section, "\n".join(current_lines)))

    # Reorganize sections to ensure Chapter (BOB) and Subchapter (1.1) titles are back-to-back
    normalized = []
    for name, text in sections_list:
        curr_name = name
        curr_lines = []
        for line in text.split('\n'):
            clean_l = line.replace("*", "").strip()
            if bool(re.match(r"^\d+\.\d+\.?\s", clean_l)) and not line.startswith("#"):
                if curr_lines or curr_name:
                    normalized.append((curr_name, "\n".join(curr_lines)))
                curr_name = clean_l
                curr_lines = []
            else:
                curr_lines.append(line)
        normalized.append((curr_name, "\n".join(curr_lines)))
        
    final_sections = []
    pending_text = ""
    for name, text in normalized:
        name_upper = name.upper()
        is_bob = "BOB." in name_upper or "BOB " in name_upper
        
        if is_bob:
            final_sections.append((name, ""))
            pending_text = text.strip()
        else:
            combined = pending_text + ("\n\n" if pending_text and text.strip() else "") + text.strip()
            pending_text = ""
            final_sections.append((name, combined))
            
    if pending_text and final_sections:
        final_sections[-1] = (final_sections[-1][0], final_sections[-1][1] + "\n\n" + pending_text)
        
    sections_list = final_sections

    # 2. Extract structured plan names
    plan_sections = [s.strip() for s in plan.split("\n") if s.strip()] if plan else []
    
    # Title Page Replacements
    for p in doc.paragraphs[:30]:
        text = p.text
        if not text.strip():
            continue
            
        modified = False
        
        if "[Universitet nomi]" in text:
            text = text.replace("[Universitet nomi]", university.upper() if university else "URGANCH DAVLAT UNIVERSITETI")
            modified = True
            
        if "[Fakultet nomi]" in text:
            if fakultet:
                text = text.replace("[Fakultet nomi]", fakultet.upper())
            else:
                text = text.replace("[Fakultet nomi]", "____________________")
            modified = True
            
        if "[Mutaxassislik]" in text:
            if specialty:
                text = text.replace("[Mutaxassislik]", specialty.upper())
            else:
                text = text.replace("[Mutaxassislik]", "____________________")
            modified = True
            
        if "[F.I.Sh]" in text:
            text = text.replace("[F.I.Sh]", author.upper() if author else "____________________")
            modified = True
            
        if "[Fan nomi]" in text:
            if subject:
                text = text.replace("[Fan nomi]", subject.upper())
            else:
                text = text.replace("[Fan nomi]", "____________________")
            modified = True
            
        if "[Mavzu]" in text:
            text = text.replace("[Mavzu]", topic.upper() if topic else "")
            modified = True
            
        if "[Bajardi]" in text:
            if author:
                text = text.replace("[Bajardi]", format_fio(author))
            else:
                text = text.replace("[Bajardi]", "____________________")
            modified = True
            
        if "[Qabul qildi]" in text:
            if reviewer:
                text = text.replace("[Qabul qildi]", reviewer)
            else:
                text = text.replace("[Qabul qildi]", "____________________")
            modified = True
            
        # Cleanup double spaces
        text = text.replace("  ", " ").strip()
            
        if modified:
            alignment = p.alignment
            
            # Save original formatting from first non-empty run if available
            is_bold = False
            is_italic = False
            color = None
            size = None
            for r in p.runs:
                if r.text.strip():
                    is_bold = r.bold
                    is_italic = r.italic
                    if r.font.color and r.font.color.rgb:
                        color = r.font.color.rgb
                    if r.font.size:
                        size = r.font.size
                    break
                    
            p.text = text
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = size if size else Pt(14)
                r.bold = is_bold
                r.italic = is_italic
                if color:
                    r.font.color.rgb = color
            p.alignment = alignment

    
    # Add page break before MUDARIJA
    for p in doc.paragraphs[:30]:
        if "MUDARIJA" in p.text.upper() or "MUNDARIJA" in p.text.upper():
            p.text = ""
            r = p.add_run()
            r.add_break(WD_BREAK.PAGE)
            r = p.add_run("MUNDARIJA")
            r.bold = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break
            
# 3. Handle Mundarija Table
    if doc.tables:
        table = doc.tables[0]
        # Clear existing rows except first one as a template
        while len(table.rows) > 0:
            table._tbl.remove(table.rows[-1]._tr)
            
        def add_row(col0_text, col1_text, col3_text):
            row = table.add_row()
            is_chapter = col0_text.strip().endswith("BOB.")
            
            if not col1_text and col0_text.strip() in ["KIRISH", "XULOSA", "FOYDALANILGAN ADABIYOTLAR"]:
                row.cells[0].merge(row.cells[1])
                full_text = f"{col0_text}\t"
                row.cells[0].text = full_text
                c1_run = row.cells[0].paragraphs[0].runs[0] if row.cells[0].paragraphs[0].runs else row.cells[0].paragraphs[0].add_run(full_text)
                c1_run.font.name = "Times New Roman"
                c1_run.font.size = Pt(14)
                
                # Add tab stop with dot leader
                tab_stops = row.cells[0].paragraphs[0].paragraph_format.tab_stops
                tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                
                row.cells[2].text = str(col3_text)
                c3_run = row.cells[2].paragraphs[0].runs[0] if row.cells[2].paragraphs[0].runs else row.cells[2].paragraphs[0].add_run(str(col3_text))
                c3_run.font.name = "Times New Roman"
                c3_run.font.size = Pt(14)
                return
                
            # Column 1
            row.cells[0].text = col0_text
            c1_run = row.cells[0].paragraphs[0].runs[0] if row.cells[0].paragraphs[0].runs else row.cells[0].paragraphs[0].add_run(col0_text)
            c1_run.font.name = "Times New Roman"
            c1_run.font.size = Pt(14)
            if is_chapter:
                c1_run.bold = True
            
            # Column 2 (title + dots)
            if is_chapter and col1_text:
                col1_text = col1_text.upper()
                
            full_col1 = f"{col1_text}\t" if col1_text else "\t"
            
            row.cells[1].text = full_col1
            c2_run = row.cells[1].paragraphs[0].runs[0] if row.cells[1].paragraphs[0].runs else row.cells[1].paragraphs[0].add_run(full_col1)
            c2_run.font.name = "Times New Roman"
            c2_run.font.size = Pt(14)
            if is_chapter:
                c2_run.bold = True
                
            # Add tab stop with dot leader
            tab_stops = row.cells[1].paragraphs[0].paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            c2_run.font.name = "Times New Roman"
            c2_run.font.size = Pt(14)
            if is_chapter:
                c2_run.bold = True
            
            # Column 3 (page)
            row.cells[2].text = str(col3_text)
            c3_run = row.cells[2].paragraphs[0].runs[0] if row.cells[2].paragraphs[0].runs else row.cells[2].paragraphs[0].add_run(str(col3_text))
            c3_run.font.name = "Times New Roman"
            c3_run.font.size = Pt(14)
            
        # Estimate page numbers dynamically based on actual line wrapping
        current_page_est = 3
        current_lines = 0
        LINES_PER_PAGE = 29
        CHARS_PER_LINE = 68
        
        section_pages = {}
        for name, text in sections_list:
            name_upper = name.upper()
            is_xulosa = "XULOSA" in name_upper
            is_adabiyot = "ADABIYOT" in name_upper or "FOYDALANILGAN" in name_upper
            is_kirish = "KIRISH" in name_upper
            is_bob = "BOB." in name_upper or "BOB " in name_upper
            
            should_page_break = is_xulosa or is_adabiyot or is_kirish or is_bob
            if should_page_break and current_lines > 0:
                current_page_est += 1
                current_lines = 0
                
            section_pages[name_upper] = current_page_est
            
            current_lines += 2 # Title spacing
            
            paragraphs = text.split('\n')
            for p in paragraphs:
                p = p.strip()
                if not p:
                    current_lines += 1
                else:
                    lines_for_p = (len(p) + CHARS_PER_LINE - 1) // CHARS_PER_LINE
                    current_lines += lines_for_p
                    
            while current_lines >= LINES_PER_PAGE:
                current_page_est += 1
                current_lines -= LINES_PER_PAGE

        def get_page(sec_name, default_page):
            sec_upper = sec_name.upper().replace("*", "").strip()
            # 1. Exact match
            for k, v in section_pages.items():
                if k == sec_upper: return v
            
            # 2. Prefix match (e.g., "1.1.", "I BOB")
            prefix_match = re.match(r'^([IVX]+[-\s]*BOB|\d+\.\d+)', sec_upper)
            if prefix_match:
                prefix = prefix_match.group(1)
                for k, v in section_pages.items():
                    if k.startswith(prefix): return v
                        
            # 3. Substring match (clean numbering)
            clean_sec = re.sub(r'^([IVX]+[-\s]*BOB\.?|\d+\.\d+\.?)\s*', '', sec_upper).strip()
            if clean_sec and len(clean_sec) > 5:
                for k, v in section_pages.items():
                    if clean_sec in k: return v
                    
            return default_page

        fallback_page = 3
        curr_pg = get_page("KIRISH", fallback_page)
        add_row("KIRISH", "", curr_pg)
        fallback_page = curr_pg
        
        chapter_idx = 1
        chapter_numerals = ["I", "II", "III", "IV", "V", "VI", "VII"]
        for sec in plan_sections:
            if sec.upper() in ["KIRISH", "XULOSA", "FOYDALANILGAN ADABIYOTLAR"]:
                continue
                
            sec_clean = sec.replace("**", "").replace("*", "").strip()
            ch_match = re.match(r'^([IVX]+|\d+)[-\s]*BOB\.?\s*(.+)', sec_clean, flags=re.IGNORECASE)
            if ch_match:
                ch_val = ch_match.group(1).upper()
                roman_map = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6}
                ch = int(ch_val) if ch_val.isdigit() else roman_map.get(ch_val, 1)
                title = ch_match.group(2)
                
                pg = get_page(sec_clean, fallback_page)
                fallback_page = pg
                
                if ch <= len(chapter_numerals):
                    add_row(f"{chapter_numerals[ch-1]} BOB.", title, pg)
                    chapter_idx = ch + 1
                continue
                
            match = re.match(r'^(\d+)\.(\d+)\.?\s*(.+)', sec_clean)
            if match:
                ch = int(match.group(1))
                sub = int(match.group(2))
                title = match.group(3)
                
                pg = get_page(sec_clean, fallback_page)
                fallback_page = pg
                
                if ch >= chapter_idx and ch <= len(chapter_numerals):
                    ch_pg = get_page(f"{chapter_numerals[ch-1]} BOB.", pg)
                    add_row(f"{chapter_numerals[ch-1]} BOB.", "", ch_pg)
                    chapter_idx = ch + 1
                add_row(f"{ch}.{sub}.", title, pg)
            else:
                pg = get_page(sec_clean, fallback_page)
                fallback_page = pg
                add_row(sec, "", pg)
                
        pg = get_page("XULOSA", fallback_page + 1)
        add_row("XULOSA", "", pg)
        fallback_page = pg
        
        pg = get_page("FOYDALANILGAN ADABIYOTLAR", fallback_page + 1)
        add_row("FOYDALANILGAN ADABIYOTLAR", "", pg)

    # 4. Remove everything after the table to rebuild dynamically
    body = doc._body._body
    found_table = False
    to_delete = []
    for child in body.iterchildren():
        if child.tag.endswith('tbl'):
            found_table = True
            continue
        if found_table and (child.tag.endswith('p') or child.tag.endswith('tbl')):
            to_delete.append(child)
            
    for elem in to_delete:
        body.remove(elem)
            
    # Add body from AI content
    def add_section_header(text, page_break_before=False):
        p = doc.add_paragraph()
        if page_break_before:
            p.paragraph_format.page_break_before = True
            
        clean_name = text.replace("*", "").strip()
        is_subchapter = bool(re.match(r"^\d+\.\d+\.?\s", clean_name))

        # Make the header text uppercase if it's not a subchapter
        display_text = text if is_subchapter else text.upper()
        
        r = p.add_run(display_text)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def add_section_text(text_block):
        lines = [l.strip() for l in text_block.split("\n") if l.strip()]
        for l in lines:
            if l.startswith("#"): continue
            
            # Detect subchapters like "1.1. Buxoro...", optionally with bold asterisks
            clean_l = l.replace("*", "").strip()
            is_subchapter = bool(re.match(r"^\d+\.\d+\..*", clean_l))
            
            new_p = doc.add_paragraph()
            if is_subchapter:
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_p.paragraph_format.first_line_indent = Cm(0)
            else:
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_p.paragraph_format.first_line_indent = Cm(1.25)
                
            new_p.paragraph_format.line_spacing = 1.5
            
            clean_parts = re.split(r"(\*\*[^*]+\*\*)", l)
            for part in clean_parts:
                if not part: continue
                run = new_p.add_run(part[2:-2] if part.startswith("**") else part)
                if part.startswith("**") or is_subchapter: 
                    run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)

    for name, text in sections_list:
        name_upper = name.upper()
        is_xulosa = "XULOSA" in name_upper
        is_adabiyot = "ADABIYOT" in name_upper or "FOYDALANILGAN" in name_upper
        is_kirish = "KIRISH" in name_upper
        is_bob = "BOB." in name_upper or "BOB " in name_upper
        
        should_page_break = is_xulosa or is_adabiyot or is_kirish or is_bob
        add_section_header(name, page_break_before=should_page_break)
        add_section_text(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def generate_maqola_from_template(
    topic: str,
    sections_content: dict,
    author: str = "Talaba",
    university: str = "",
    plan_titles: dict = None,
    images: dict = None,
) -> bytes:
    """
    Generate a maqola DOCX using Maqola_template.docx.
    Template has: title page (Mavzu + table with author info).
    All sections are appended programmatically after the title page.
    images: dict mapping placeholder text -> image bytes
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    template_path = os.path.join(TEMPLATE_DIR, "Maqola_template.docx")
    doc = Document(template_path)
    plan_titles = plan_titles or {}
    images = images or {}
    generate_maqola_from_template._img_counter = 0

    # ── 1. Fill title placeholder ─────────────────────────────────────────
    for p in doc.paragraphs[:5]:
        if "___" in p.text or "Mavzu" in p.text:
            if p.runs:
                p.runs[0].text = topic.upper()
                p.runs[0].font.name = "Times New Roman"
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
                for run in p.runs[1:]:
                    run.text = ""
            break

    # ── 2. Fill author table ──────────────────────────────────────────────
    if doc.tables:
        table = doc.tables[0]
        # Cell [0,1] = author name + university
        if len(table.rows) > 0 and len(table.columns) > 1:
            cell = table.cell(0, 1)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ""
            # Set author name (bold italic)
            first_p = cell.paragraphs[0]
            if first_p.runs:
                first_p.runs[0].text = author
            else:
                run = first_p.add_run(author)
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                run.bold = True
                run.italic = True
            # Add university line
            if university and len(cell.paragraphs) > 1:
                second_p = cell.paragraphs[1]
                if second_p.runs:
                    second_p.runs[0].text = university
                    for run in second_p.runs[1:]:
                        run.text = ""
                else:
                    run = second_p.add_run(university)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)
                    run.italic = True
            elif university:
                up = cell.add_paragraph()
                run = up.add_run(university)
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                run.italic = True

    # ── 3. Remove empty paragraphs after table ────────────────────────────
    for p in list(doc.paragraphs):
        if not p.text.strip() and p._element.getparent() is not None:
            parent = p._element.getparent()
            parent.remove(p._element)

    # ── 4. Build junk line filter ─────────────────────────────────────────
    topic_upper = topic.upper().strip()
    skip_phrases = {
        "KIRISH", "XULOSA", "ANNOTATSIYA", "KALIT SO'ZLAR",
        "FOYDALANILGAN ADABIYOTLAR", "FOYDALANILGAN ADABIYOTLAR RO'YXATI",
        "ASOSIY QISM", "XULOSA VA TAKLIFLAR", "MUHOKAMA VA NATIJALAR",
        "ADABIYOTLAR TAHLILI", "METODOLOGIYA", "ABSTRACT", "KEYWORDS",
    }
    for title in plan_titles.values():
        skip_phrases.add(title.upper().strip())

    def _is_junk(line: str, section_name: str) -> bool:
        lu = line.upper().strip().rstrip(":")
        if not lu:
            return True
        if lu.startswith("MAVZU:") or lu.startswith("MAVZU "):
            return True
        if lu == section_name.upper().rstrip(":"):
            return True
        if lu in skip_phrases:
            return True
        stripped = re.sub(r'^\d+(?:\.\d+)?\.\s*', '', lu)
        if stripped and stripped in skip_phrases:
            return True
        if lu == topic_upper:
            return True
        # Short ALL-CAPS line matching topic words
        if len(lu) < 80 and line.strip().isupper():
            tw = set(topic_upper.split())
            lw = set(lu.split())
            if tw and len(tw & lw) >= len(tw) * 0.5:
                return True
        return False

    # ── 5. Helper: create paragraph with Times New Roman 14pt ─────────────
    def _make_para(text, bold=False, center=False, first_indent=True):
        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        if center:
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
        else:
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'both')
            pPr.append(jc)
        if first_indent and not center:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLine'), '709')
            pPr.append(ind)
        # Line spacing 1.15
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '276')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        new_p.append(pPr)

        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '28')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '28')
        rPr.append(szCs)
        if bold:
            b = OxmlElement('w:b'); rPr.append(b)
        run_elem.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        run_elem.append(t)
        new_p.append(run_elem)
        return new_p

    # ── 6. Define section order ───────────────────────────────────────────
    section_order = [
        ("annotatsiya", "ANNOTATSIYA:"),
        ("kalit_sozlar", "KALIT SO'ZLAR:"),
        ("kirish", "KIRISH"),
        ("1", f"1. {plan_titles.get('1', '')}"),
        ("1.1", f"1.1. {plan_titles.get('1.1', '')}"),
        ("1.2", f"1.2. {plan_titles.get('1.2', '')}"),
        ("2", f"2. {plan_titles.get('2', '')}"),
        ("2.1", f"2.1. {plan_titles.get('2.1', '')}"),
        ("2.2", f"2.2. {plan_titles.get('2.2', '')}"),
        ("xulosa", "XULOSA"),
        ("adabiyotlar", "FOYDALANILGAN ADABIYOTLAR"),
    ]

    # ── 7. Append sections to document body ───────────────────────────────
    body = doc.element.body
    # Find sectPr (must stay last)
    sect_pr = body.find(qn('w:sectPr'))

    for content_key, section_title in section_order:
        ai_text = sections_content.get(content_key, "")

        # Add section header (bold, centered)
        header_p = _make_para(section_title, bold=True, center=True, first_indent=False)
        if sect_pr is not None:
            sect_pr.addprevious(header_p)
        else:
            body.append(header_p)

        if not ai_text:
            continue

        # Filter and add content lines
        text_lines = [l.strip() for l in ai_text.strip().split("\n") if l.strip()]
        img_counter = getattr(generate_maqola_from_template, '_img_counter', 0)
        for line in text_lines:
            if line.startswith("#"):
                continue
            if _is_junk(line, section_title):
                continue
            # Check if line is an AI image placeholder [ RASM: ... ]
            # ([ 🖼️ SHU YERGA RASM JOYLANG: ... ] is kept as-is for manual insertion)
            img_match = re.search(r'\[\s*RASM:\s*(.+?)\s*\]', line)
            if img_match and 'SHU YERGA' not in line and 'JOYLANG' not in line:
                placeholder_text = img_match.group(0)
                img_bytes = images.get(placeholder_text) if images else None
                if img_bytes:
                    img_counter += 1
                    # Insert image
                    img_stream = io.BytesIO(img_bytes)
                    img_p = OxmlElement('w:p')
                    img_pPr = OxmlElement('w:pPr')
                    img_jc = OxmlElement('w:jc')
                    img_jc.set(qn('w:val'), 'center')
                    img_pPr.append(img_jc)
                    img_p.append(img_pPr)
                    if sect_pr is not None:
                        sect_pr.addprevious(img_p)
                    else:
                        body.append(img_p)
                    # Use python-docx to add the picture via a temporary paragraph
                    # We need to get the paragraph object from the element
                    from docx.text.paragraph import Paragraph
                    para_obj = Paragraph(img_p, doc)
                    run = para_obj.add_run()
                    run.add_picture(img_stream, width=Inches(4.5))
                    # Add caption below image
                    desc = img_match.group(1).strip()
                    caption_p = _make_para(
                        f"Rasm {img_counter}. {desc}",
                        bold=False, center=True, first_indent=False
                    )
                    if sect_pr is not None:
                        sect_pr.addprevious(caption_p)
                    else:
                        body.append(caption_p)
                    continue
                else:
                    # No image generated for this placeholder — skip the line
                    continue
            content_p = _make_para(line, bold=False, center=False, first_indent=True)
            if sect_pr is not None:
                sect_pr.addprevious(content_p)
            else:
                body.append(content_p)
        generate_maqola_from_template._img_counter = img_counter

    buf_out = io.BytesIO()
    doc.save(buf_out)
    buf_out.seek(0)
    return buf_out.read()
