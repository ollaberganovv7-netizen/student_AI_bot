from __future__ import annotations
"""
Avtomatik Quiz Tuzish handler.
Architecture:
1. File upload → text extraction (DOCX/DOC/TXT/PDF)
2. AI normalization & auto-correction of format
3. Parser: find +++ blocks, validate structure
4. Validation report with per-question errors
5. Quiz name → count → time → shuffle → summary
6. Save to DB → send as Telegram Quiz polls
"""

import io
import re
import json
import random
import logging
import math
import asyncio
import time as _time

from aiogram import Router, F
from aiogram.filters import Command
from aiogram.types import (
    Message, CallbackQuery, PollAnswer,
    InlineKeyboardMarkup, InlineKeyboardButton,
    ReplyKeyboardMarkup, KeyboardButton,
)
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

from database.models import User
from database.db import save_quiz, get_user_quizzes, get_quiz_with_questions, delete_quiz
from keyboards.main_kb import main_menu_kb
from utils.i18n import btn

logger = logging.getLogger(__name__)

router = Router()


class QuizStates(StatesGroup):
    waiting_for_file = State()
    entering_name = State()
    choosing_count = State()
    entering_custom_count = State()
    choosing_time = State()
    choosing_shuffle = State()
    confirming = State()
    playing = State()  # Sequential quiz play


def _lang(db_user) -> str:
    return (db_user.language or "uz") if db_user else "uz"


# ─── Quiz menu keyboard ──────────────────────────────────────────────────────

def quiz_menu_kb() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text=" Avtomatik quiz tuzish")],
            [KeyboardButton(text="📂 Mening quizlarim")],
            [KeyboardButton(text="🏠 Bosh menyu")],
        ],
        resize_keyboard=True,
    )


# ══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

async def _extract_text_from_file(message: Message) -> tuple[str | None, str]:
    """Download and extract text. Returns (text, format_name)."""
    doc = message.document
    if not doc:
        return None, ""
    
    file_name = doc.file_name or ""
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    
    file = await message.bot.download(doc)
    file_bytes = file.read()
    
    if ext == "txt":
        try:
            return file_bytes.decode("utf-8"), "TXT"
        except UnicodeDecodeError:
            return file_bytes.decode("cp1251", errors="replace"), "TXT"
    
    elif ext in ("docx", "doc"):
        # Try python-docx first (works for .docx)
        try:
            from docx import Document as DocxDocument
            doc_obj = DocxDocument(io.BytesIO(file_bytes))
            
            # Simple paragraph text extraction
            lines = [p.text for p in doc_obj.paragraphs]
            
            # Also extract from tables
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            lines.append(cell_text)
            
            text = "\n".join(lines)
            logger.info(f"DOCX extracted: {len(text)} chars")
            return text, ext.upper()
        except Exception as e:
            logger.warning(f"python-docx failed: {e}, trying fallback...")
        
        # Fallback: save to temp file and use textutil (macOS) or strings
        import tempfile
        import subprocess
        try:
            with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            
            # Try textutil (macOS built-in)
            try:
                txt_path = tmp_path + ".txt"
                subprocess.run(
                    ["textutil", "-convert", "txt", "-output", txt_path, tmp_path],
                    capture_output=True, timeout=15
                )
                import os
                if os.path.exists(txt_path):
                    with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
                        result = f.read()
                    os.unlink(txt_path)
                    os.unlink(tmp_path)
                    if result.strip():
                        return result, ext.upper()
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
            
            # Last fallback: extract readable strings from binary
            import os
            os.unlink(tmp_path)
            text = file_bytes.decode("utf-8", errors="ignore")
            # Filter only printable lines
            lines = [l for l in text.split("\n") if l.strip() and any(c.isalpha() for c in l)]
            if lines:
                return "\n".join(lines), ext.upper()
        except Exception as e2:
            logger.error(f"DOC fallback error: {e2}")
        
        return None, ext.upper()
    
    elif ext == "pdf":
        try:
            import fitz
            pdf = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            for page in pdf:
                text += page.get_text()
            pdf.close()
            return text, "PDF"
        except Exception as e:
            logger.error(f"PDF parse error: {e}")
            return None, "PDF"
    
    return None, ext.upper()


# ══════════════════════════════════════════════════════════════════════════════
# AI NORMALIZATION
# ══════════════════════════════════════════════════════════════════════════════

def _normalize_text(text: str) -> str:
    """Code-based normalization before AI step."""
    # Remove invisible chars
    text = text.replace('\u200b', '').replace('\ufeff', '').replace('\xa0', ' ')
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # Remove excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Fix delimiters ONLY when they appear on their own line
    text = re.sub(r'^\s*\+{2,}\s*$', '+++', text, flags=re.MULTILINE)   # ++/++++ → +++
    text = re.sub(r'^\s*={2,}\s*$', '===', text, flags=re.MULTILINE)    # ==/==== → ===
    # Fix # with extra spaces only at line start (answer marker)
    text = re.sub(r'^\s*#\s{2,}', '#', text, flags=re.MULTILINE)        # #   javob → #javob
    text = re.sub(r'^\s*#\s*#\s*', '#', text, flags=re.MULTILINE)       # # # → #
    return text.strip()


async def _ai_normalize(text: str) -> str:
    """Use Gemini AI to fix formatting errors in quiz text.
    Only applies to small texts (<10000 chars). For large files,
    code normalization is sufficient."""
    
    # Skip AI for large files — code normalization is enough
    if len(text) > 10000:
        logger.info(f"Skipping AI normalize for large text ({len(text)} chars)")
        return text
    
    from services.gemini_service import _call_gemini, GeminiError
    
    prompt = (
        "Sen quiz fayl formatlash yordamchisisiz.\n"
        "Quyidagi matnni tekshir va formatlash xatolarini tuzat.\n\n"
        "QOIDALAR:\n"
        "- Har bir savol +++ bilan boshlanib +++ bilan tugashi kerak\n"
        "- Variantlar === bilan ajratilishi kerak\n"
        "- To'g'ri javob oldida # belgisi bo'lishi kerak\n"
        "- Har savolda faqat 1 ta # bo'lishi kerak\n"
        "- == bo'lsa === ga tuzat\n"
        "- ++ bo'lsa +++ ga tuzat\n"
        "- # # javob bo'lsa #javob ga tuzat\n"
        "- Bo'sh variantlarni o'chir\n"
        "- Savol matni bo'sh bo'lsa o'chirish kerak\n\n"
        "MUHIM: Faqat tuzatilgan matnni qaytar, boshqa hech narsa yozma.\n"
        "Agar matn to'g'ri bo'lsa, o'zgartirmasdan qaytar.\n\n"
        f"MATN:\n{text}"
    )
    
    try:
        result = await _call_gemini(prompt, use_search=False, temperature=0.1, max_tokens=8000)
        if result and '+++' in result and '===' in result:
            return result
    except (GeminiError, Exception) as e:
        logger.warning(f"Gemini normalize failed: {e}")
    
    return text  # fallback to original


# ══════════════════════════════════════════════════════════════════════════════
# PARSER + VALIDATOR
# ══════════════════════════════════════════════════════════════════════════════

def _parse_quiz_text(text: str) -> tuple[list[dict], list[str]]:
    """
    Parse and validate quiz text.
    Returns (valid_questions, error_messages).
    """
    questions = []
    errors = []
    blocks = re.split(r'\+{3,}', text)
    
    block_num = 0
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        block_num += 1
        
        parts = re.split(r'={3,}', block)
        parts = [p.strip() for p in parts if p.strip()]
        
        # Validation
        if len(parts) < 2:
            errors.append(f"❌ {block_num}-savol: Variantlar topilmadi")
            continue
        
        question_text = parts[0]
        if not question_text:
            errors.append(f"❌ {block_num}-savol: Savol matni bo'sh")
            continue
        
        options = []
        correct_index = None
        correct_count = 0
        
        for part in parts[1:]:
            stripped = part.lstrip()
            if stripped.startswith('#'):
                correct_count += 1
                correct_index = len(options)
                options.append(stripped[1:].strip())
            else:
                options.append(part.strip())
        
        # Remove empty options
        clean_options = []
        new_correct = None
        for i, opt in enumerate(options):
            if opt:
                if i == correct_index:
                    new_correct = len(clean_options)
                clean_options.append(opt)
        options = clean_options
        correct_index = new_correct
        
        if correct_count == 0:
            errors.append(f"❌ {block_num}-savol: To'g'ri javob topilmadi (#)")
            continue
        if correct_count > 1:
            errors.append(f"❌ {block_num}-savol: {correct_count} ta to'g'ri javob bor (faqat 1 ta bo'lishi kerak)")
            continue
        if len(options) < 2:
            errors.append(f"❌ {block_num}-savol: Variantlar soni yetarli emas ({len(options)} ta)")
            continue
        if len(options) > 10:
            errors.append(f"⚠️ {block_num}-savol: {len(options)} ta variant (max 10), ortiqchasi kesildi")
            options = options[:10]
            if correct_index >= 10:
                errors.append(f"❌ {block_num}-savol: To'g'ri javob kesilgan variantlar orasida")
                continue
        
        questions.append({
            "question": question_text[:300],
            "options": [o[:100] for o in options],
            "correct_index": correct_index,
        })
    
    return questions, errors


def _build_report(questions: list, errors: list, fmt: str) -> str:
    """Build validation report."""
    report = f"✅ <b>Fayl muvaffaqiyatli o'qildi!</b>\n\n"
    report += f"🖥 <b>Format:</b> {fmt}\n"
    report += f"❓ <b>Savollar soni:</b> {len(questions)} ta\n"
    
    if errors:
        report += f"\n⚠️ <b>{len(errors)} ta xatolik topildi:</b>\n\n"
        for err in errors[:10]:  # Show max 10 errors
            report += f"{err}\n"
        if len(errors) > 10:
            report += f"\n<i>... va yana {len(errors) - 10} ta xatolik</i>\n"
    
    return report


def _build_summary(data: dict) -> str:
    name = data.get("quiz_name", "Quiz")
    total = data.get("total_questions", 0)
    per_quiz = data.get("per_quiz", total)
    time_sec = data.get("time_per_q", 30)
    shuffle = data.get("shuffle_mode", "none")
    num_quizzes = math.ceil(total / per_quiz) if per_quiz else 1
    
    shuffle_label = {
        "none": "Aralashtirilmasin",
        "options": "Faqat variantlar",
        "all": "Barchasi",
    }.get(shuffle, shuffle)
    
    return (
        f"📋 <b>Quiz sozlamalari</b>\n\n"
        f"📌 <b>Nom:</b> {name}\n"
        f"❓ <b>Savollar:</b> {total} ta\n"
        f"📝 <b>Har quiz:</b> {per_quiz} ta savol\n"
        f"⏱ <b>Har savol uchun:</b> {time_sec} soniya\n"
        f"🔀 <b>Aralashtirish:</b> {shuffle_label}\n"
        f"📊 <b>Quizlar soni:</b> {num_quizzes} ta"
    )


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1: Entry — show instructions & wait for file
# ══════════════════════════════════════════════════════════════════════════════

@router.message(F.text == "📋 Avtomatik quiz tuzish")
async def start_quiz(message: Message, state: FSMContext, db_user: User):
    await state.clear()
    
    instruction = (
        "📎 <b>Fayl yuboring</b>\n\n"
        "🖥 <b>Qo'llab-quvvatlanadigan formatlar:</b> DOCX, DOC, TXT, PDF\n\n"
        "📝 <b>Fayl formati:</b>\n"
        "Har bir savol quyidagi shaklda bo'lishi kerak:\n\n"
        "<code>+++\n"
        "Savol matni?\n"
        "===\n"
        "#To'g'ri javob\n"
        "===\n"
        "Noto'g'ri javob 1\n"
        "===\n"
        "Noto'g'ri javob 2\n"
        "===\n"
        "Noto'g'ri javob 3\n"
        "+++</code>\n\n"
        "⚠️ <b>Muhim:</b>\n"
        "• <code>+++</code> — savol boshlanishi va tugashi\n"
        "• <code>===</code> — variantlar orasidagi ajratuvchi\n"
        "• <code>#</code> — to'g'ri javob oldiga qo'yiladi\n"
        "• Har bir savolda faqat <b>bitta</b> to'g'ri javob bo'lishi kerak\n"
        "• Variantlar soni 2 dan 10 gacha bo'lishi mumkin"
    )
    
    await message.answer(instruction, parse_mode="HTML", reply_markup=quiz_menu_kb())
    await state.set_state(QuizStates.waiting_for_file)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2: File received → normalize → parse → validate → report → ask name
# ══════════════════════════════════════════════════════════════════════════════

@router.message(QuizStates.waiting_for_file, F.document)
async def process_quiz_file(message: Message, state: FSMContext, db_user: User):
    file_name = message.document.file_name or ""
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    
    if ext not in ("txt", "docx", "doc", "pdf"):
        await message.answer(
            "⚠️ <b>Noto'g'ri format!</b>\n\nFaqat DOCX, DOC, TXT yoki PDF fayl yuboring.",
            parse_mode="HTML"
        )
        return
    
    wait_msg = await message.answer("⏳ <b>Fayl tahlil qilinmoqda...</b>", parse_mode="HTML")
    
    # 1. Extract text
    text, fmt = await _extract_text_from_file(message)
    if not text:
        await wait_msg.edit_text(
            "❌ <b>Faylni o'qib bo'lmadi!</b>\n\nIltimos, boshqa formatda yuboring.",
            parse_mode="HTML"
        )
        return
    
    # 2. Code normalization
    try:
        await wait_msg.edit_text("🔧 <b>Matn normalizatsiya qilinmoqda...</b>", parse_mode="HTML")
    except TelegramBadRequest:
        pass
    logger.info(f"Raw text preview (500 chars): {text[:500]!r}")
    logger.info(f"Raw markers: +++ count={text.count('+++')}, === count={text.count('===')}, # count={text.count('#')}")
    text = _normalize_text(text)
    
    # 3. AI normalization (auto-fix errors)
    try:
        await wait_msg.edit_text("🤖 <b>AI xatolarni tekshirmoqda...</b>", parse_mode="HTML")
    except TelegramBadRequest:
        pass
    text = await _ai_normalize(text)
    
    # 4. Parse + validate
    questions, errors = _parse_quiz_text(text)
    
    if not questions:
        error_text = "❌ <b>Savollar topilmadi!</b>\n\n"
        if errors:
            error_text += "📋 <b>Xatolar:</b>\n"
            for err in errors[:15]:
                error_text += f"{err}\n"
        else:
            error_text += (
                "Fayl formatini tekshiring:\n"
                "• <code>+++</code> bilan savolni boshlang/tugating\n"
                "• <code>===</code> bilan variantlarni ajrating\n"
                "• <code>#</code> bilan to'g'ri javobni belgilang"
            )
        await wait_msg.edit_text(error_text, parse_mode="HTML")
        return
    
    # 5. Show report
    report = _build_report(questions, errors, fmt)
    report += "\n\nEndi quiz nomini kiriting:"
    await wait_msg.edit_text(report, parse_mode="HTML")
    
    await state.update_data(
        questions=questions,
        total_questions=len(questions),
        file_format=fmt,
    )
    await state.set_state(QuizStates.entering_name)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3: Name → count
# ══════════════════════════════════════════════════════════════════════════════

@router.message(QuizStates.entering_name, F.text)
async def set_quiz_name(message: Message, state: FSMContext, db_user: User):
    name = message.text.strip()
    
    # Handle exit commands
    if name in ("🏠 Bosh menyu", "/stop", "/start", "/cancel"):
        await state.clear()
        lang = _lang(db_user)
        await message.answer("✅ Bekor qilindi.", reply_markup=main_menu_kb(lang))
        return
    
    if not name or name.startswith("/") or name in ("📋 Avtomatik quiz tuzish", "📂 Mening quizlarim"):
        await message.answer("📌 Iltimos, quiz nomini kiriting:")
        return
    
    await state.update_data(quiz_name=name)
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="10", callback_data="qcount:10"),
            InlineKeyboardButton(text="15", callback_data="qcount:15"),
            InlineKeyboardButton(text="20", callback_data="qcount:20"),
        ],
        [
            InlineKeyboardButton(text="25", callback_data="qcount:25"),
            InlineKeyboardButton(text="30", callback_data="qcount:30"),
            InlineKeyboardButton(text="40", callback_data="qcount:40"),
        ],
        [
            InlineKeyboardButton(text="50", callback_data="qcount:50"),
            InlineKeyboardButton(text="Boshqa", callback_data="qcount:custom"),
        ],
    ])
    
    await message.answer("Har quizda nechta savol bo'lsin?", reply_markup=kb)
    await state.set_state(QuizStates.choosing_count)


@router.callback_query(F.data.startswith("qcount:"), QuizStates.choosing_count)
async def set_quiz_count(callback: CallbackQuery, state: FSMContext, db_user: User):
    val = callback.data.split(":")[1]
    await callback.message.delete()
    
    if val == "custom":
        await callback.message.answer("Boshqa son kiriting (2..200).")
        await state.set_state(QuizStates.entering_custom_count)
        await callback.answer()
        return
    
    count = int(val)
    await state.update_data(per_quiz=count)
    await callback.answer()
    await _ask_time(callback.message, state)


@router.message(QuizStates.entering_custom_count, F.text)
async def set_custom_count(message: Message, state: FSMContext, db_user: User):
    try:
        count = int(message.text.strip())
        if count < 2 or count > 200:
            raise ValueError
    except ValueError:
        await message.answer("⚠️ 2 dan 200 gacha son kiriting.")
        return
    
    await state.update_data(per_quiz=count)
    await _ask_time(message, state)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 4: Time per question
# ══════════════════════════════════════════════════════════════════════════════

async def _ask_time(message: Message, state: FSMContext):
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="15", callback_data="qtime:15"),
            InlineKeyboardButton(text="30", callback_data="qtime:30"),
        ],
        [
            InlineKeyboardButton(text="60", callback_data="qtime:60"),
            InlineKeyboardButton(text="120", callback_data="qtime:120"),
        ],
    ])
    await message.answer("Har savol uchun vaqtni tanlang.", reply_markup=kb)
    await state.set_state(QuizStates.choosing_time)


@router.callback_query(F.data.startswith("qtime:"), QuizStates.choosing_time)
async def set_quiz_time(callback: CallbackQuery, state: FSMContext, db_user: User):
    time_sec = int(callback.data.split(":")[1])
    await state.update_data(time_per_q=time_sec)
    await callback.message.delete()
    await callback.answer()
    await _ask_shuffle(callback.message, state)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 5: Shuffle mode
# ══════════════════════════════════════════════════════════════════════════════

async def _ask_shuffle(message: Message, state: FSMContext):
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Aralashtirilmasin", callback_data="qshuf:none")],
        [InlineKeyboardButton(text="Faqat variantlar", callback_data="qshuf:options")],
        [InlineKeyboardButton(text="Barchasini aralashtirish", callback_data="qshuf:all")],
    ])
    await message.answer("Savollar va javob variantlari aralashtirilsinmi?", reply_markup=kb)
    await state.set_state(QuizStates.choosing_shuffle)


@router.callback_query(F.data.startswith("qshuf:"), QuizStates.choosing_shuffle)
async def set_shuffle(callback: CallbackQuery, state: FSMContext, db_user: User):
    mode = callback.data.split(":")[1]
    await state.update_data(shuffle_mode=mode)
    await callback.message.delete()
    await callback.answer()
    
    data = await state.get_data()
    summary = _build_summary(data)
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Yaratish", callback_data="quiz_create")],
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="main_menu")],
    ])
    
    await callback.message.answer(summary, parse_mode="HTML", reply_markup=kb)
    await state.set_state(QuizStates.confirming)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 6: Create quiz → save to DB → show quiz cards
# ══════════════════════════════════════════════════════════════════════════════

@router.callback_query(F.data == "quiz_create", QuizStates.confirming)
async def create_quiz(callback: CallbackQuery, state: FSMContext, db_user: User):
    await callback.message.delete()
    data = await state.get_data()
    
    questions = data.get("questions", [])
    name = data.get("quiz_name", "Quiz")
    per_quiz = data.get("per_quiz", len(questions))
    time_sec = data.get("time_per_q", 30)
    shuffle_mode = data.get("shuffle_mode", "none")
    
    # Save to DB
    quiz = await save_quiz(
        user_id=db_user.id,
        name=name,
        questions=questions,
        per_quiz=per_quiz,
        time_per_q=time_sec,
        shuffle_mode=shuffle_mode,
    )
    
    num_quizzes = math.ceil(len(questions) / per_quiz)
    
    # Show quiz cards
    for qi in range(1, num_quizzes + 1):
        start_idx = (qi - 1) * per_quiz
        end_idx = min(qi * per_quiz, len(questions))
        q_count = end_idx - start_idx
        quiz_label = f"{name} {qi}" if num_quizzes > 1 else name
        
        await callback.message.answer(
            f"📋 <b>Quiz nomi:</b> {quiz_label}\n"
            f"❓ <b>Savollar:</b> {q_count} ta",
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="▶️ Boshlash / Qayta urinish", callback_data=f"qplay:{quiz.id}:{qi}")],
                [InlineKeyboardButton(text="📤 Testni ulashish", switch_inline_query=f"Quiz: {quiz_label}")],
                [InlineKeyboardButton(text="👥 Guruhda testni boshlash", callback_data=f"qgroup:{quiz.id}:{qi}")],
            ])
        )
    
    await callback.message.answer(
        f"✅ <b>{num_quizzes} ta quiz tayyor!</b>\n\n"
        f"▶️ <i>Boshlash tugmasini bosing</i>",
        parse_mode="HTML",
        reply_markup=quiz_menu_kb()
    )
    await state.clear()
    await callback.answer()


# ══════════════════════════════════════════════════════════════════════════════
# PLAY QUIZ — sequential with time tracking & group leaderboard
# ══════════════════════════════════════════════════════════════════════════════

# Track active quiz sessions: {user_id: asyncio.Event}
_quiz_answer_events: dict[int, asyncio.Event] = {}

# Group quiz sessions: {chat_id: {poll_id: session_data}}
# session_data = {"quiz_id", "questions", "index", "time_per_q", "shuffle", "participants": {user_id: {...}}}
_group_sessions: dict[int, dict] = {}


def _format_time(seconds: float) -> str:
    """Format seconds to mm:ss."""
    m = int(seconds) // 60
    s = int(seconds) % 60
    return f"{m}:{s:02d}"


def _build_leaderboard(participants: dict) -> str:
    """Build leaderboard text from participants data."""
    # Sort: score desc, then time asc
    sorted_p = sorted(
        participants.items(),
        key=lambda x: (-x[1]["correct"], x[1]["total_time"])
    )
    
    medals = ["🥇", "🥈", "🥉"]
    lines = []
    for i, (uid, data) in enumerate(sorted_p):
        medal = medals[i] if i < 3 else f"{i+1}."
        name = data.get("name", "Foydalanuvchi")
        correct = data["correct"]
        total = data["total"]
        pct = round(correct / total * 100) if total > 0 else 0
        t = _format_time(data["total_time"])
        lines.append(f"{medal} <b>{name}</b> — {correct}/{total} ({pct}%) ⏱ {t}")
    
    return "\n".join(lines) if lines else "Hech kim qatnashmadi"


async def _send_next_question(bot, chat_id: int, user_id: int, state: FSMContext):
    """Send the next question in the sequence (private chat)."""
    data = await state.get_data()
    play_questions = data.get("play_questions", [])
    play_index = data.get("play_index", 0)
    play_correct = data.get("play_correct", 0)
    time_per_q = data.get("play_time", 30)
    shuffle_mode = data.get("play_shuffle", "none")
    start_time = data.get("play_start_time", _time.time())
    total = len(play_questions)
    
    if play_index >= total:
        # Quiz finished — show detailed results
        elapsed = _time.time() - start_time
        pct = round(play_correct / total * 100) if total > 0 else 0
        
        # Grade emoji
        if pct >= 90:
            grade = "🏆 A'lo!"
        elif pct >= 70:
            grade = "👍 Yaxshi!"
        elif pct >= 50:
            grade = "📖 O'rtacha"
        else:
            grade = "📚 Ko'proq o'qing"
        
        await bot.send_message(
            chat_id,
            f"🏁 <b>Quiz tugadi!</b>\n\n"
            f"┌─────────────────────\n"
            f"│ ✅ To'g'ri: <b>{play_correct}</b> ta\n"
            f"│ ❌ Noto'g'ri: <b>{total - play_correct}</b> ta\n"
            f"│ 📊 Natija: <b>{pct}%</b>\n"
            f"│ ⏱ Vaqt: <b>{_format_time(elapsed)}</b>\n"
            f"└─────────────────────\n\n"
            f"{grade}",
            parse_mode="HTML",
            reply_markup=quiz_menu_kb()
        )
        await state.clear()
        _quiz_answer_events.pop(user_id, None)
        return
    
    q = play_questions[play_index]
    options = list(q["options"])
    correct_idx = q["correct_index"]
    
    # Shuffle options if needed
    if shuffle_mode in ("options", "all"):
        correct_answer = options[correct_idx]
        random.shuffle(options)
        correct_idx = options.index(correct_answer)
    
    await state.update_data(play_current_correct=correct_idx, play_q_start=_time.time())
    
    try:
        poll_msg = await bot.send_poll(
            chat_id,
            question=f"{play_index + 1}/{total}. {q['question'][:280]}",
            options=[opt[:100] for opt in options[:10]],
            type="quiz",
            correct_option_id=correct_idx,
            is_anonymous=False,
            open_period=time_per_q if time_per_q <= 600 else None,
        )
        await state.update_data(play_poll_id=poll_msg.poll.id)
    except Exception as e:
        logger.error(f"Quiz poll send error: {e}")
        await state.update_data(play_index=play_index + 1)
        await _send_next_question(bot, chat_id, user_id, state)
        return
    
    # Wait for answer or timeout
    event = asyncio.Event()
    _quiz_answer_events[user_id] = event
    
    try:
        await asyncio.wait_for(event.wait(), timeout=time_per_q + 2)
    except asyncio.TimeoutError:
        pass
    
    _quiz_answer_events.pop(user_id, None)
    
    current_state = await state.get_state()
    if current_state != QuizStates.playing.state:
        return
    
    data = await state.get_data()
    play_index = data.get("play_index", 0)
    await state.update_data(play_index=play_index + 1)
    
    await asyncio.sleep(0.5)
    await _send_next_question(bot, chat_id, user_id, state)


# ─── Group quiz: send next question ─────────────────────────────────────────

async def _send_group_next(bot, chat_id: int):
    """Send the next question in a group quiz session."""
    session = _group_sessions.get(chat_id)
    if not session:
        return
    
    questions = session["questions"]
    index = session["index"]
    time_per_q = session["time_per_q"]
    total = len(questions)
    
    if index >= total:
        # Quiz finished — show leaderboard
        leaderboard = _build_leaderboard(session["participants"])
        total_q = session["total"]
        
        await bot.send_message(
            chat_id,
            f"🏁 <b>Quiz tugadi!</b>\n\n"
            f"📊 <b>Natijalar ({total_q} ta savol):</b>\n\n"
            f"{leaderboard}\n\n"
            f"👥 Qatnashchilar: <b>{len(session['participants'])}</b> kishi",
            parse_mode="HTML"
        )
        _group_sessions.pop(chat_id, None)
        return
    
    q = questions[index]
    options = list(q["options"])
    correct_idx = q["correct_index"]
    
    if session["shuffle"] in ("options", "all"):
        correct_answer = options[correct_idx]
        random.shuffle(options)
        correct_idx = options.index(correct_answer)
    
    session["current_correct"] = correct_idx
    session["q_start_time"] = _time.time()
    session["answered_users"] = set()
    
    try:
        poll_msg = await bot.send_poll(
            chat_id,
            question=f"{index + 1}/{total}. {q['question'][:280]}",
            options=[opt[:100] for opt in options[:10]],
            type="quiz",
            correct_option_id=correct_idx,
            is_anonymous=False,
            open_period=time_per_q if time_per_q <= 600 else None,
        )
        session["current_poll_id"] = poll_msg.poll.id
    except Exception as e:
        logger.error(f"Group quiz poll error: {e}")
        session["index"] += 1
        await _send_group_next(bot, chat_id)
        return
    
    # Wait for timeout then move to next
    await asyncio.sleep(time_per_q + 1)
    
    # Auto-advance to next question
    if chat_id in _group_sessions and _group_sessions[chat_id]["index"] == index:
        session["index"] += 1
        await asyncio.sleep(1)
        await _send_group_next(bot, chat_id)


@router.callback_query(F.data.startswith("qplay:"))
async def play_quiz(callback: CallbackQuery, state: FSMContext, db_user: User):
    parts = callback.data.split(":")
    quiz_id = int(parts[1])
    part_num = int(parts[2]) if len(parts) > 2 else 1
    
    quiz = await get_quiz_with_questions(quiz_id)
    if not quiz:
        await callback.answer("❌ Quiz topilmadi!", show_alert=True)
        return
    
    per_quiz = quiz.per_quiz
    start_idx = (part_num - 1) * per_quiz
    end_idx = min(part_num * per_quiz, len(quiz.questions))
    
    sorted_qs = sorted(quiz.questions, key=lambda q: q.order_num)
    chunk = sorted_qs[start_idx:end_idx]
    
    if not chunk:
        await callback.answer("❌ Savollar topilmadi!", show_alert=True)
        return
    
    play_questions = [
        {"question": qq.question, "options": list(qq.options), "correct_index": qq.correct_index}
        for qq in chunk
    ]
    
    if quiz.shuffle_mode == "all":
        random.shuffle(play_questions)
    
    await state.update_data(
        play_questions=play_questions,
        play_index=0,
        play_correct=0,
        play_time=quiz.time_per_q,
        play_shuffle=quiz.shuffle_mode,
        play_quiz_id=quiz.id,
        play_start_time=_time.time(),
    )
    await state.set_state(QuizStates.playing)
    
    await callback.answer(f"▶️ {len(chunk)} ta savol boshlanmoqda...")
    await callback.message.answer(
        f"📝 <b>Quiz boshlandi!</b>\n"
        f"❓ Savollar: {len(chunk)} ta | ⏱ Vaqt: {quiz.time_per_q} soniya\n\n"
        f"<i>Har bir savolga javob bering yoki vaqt tugashini kuting.</i>",
        parse_mode="HTML"
    )
    
    await asyncio.sleep(1)
    await _send_next_question(callback.bot, callback.message.chat.id, db_user.id, state)


# ─── Group play handler ──────────────────────────────────────────────────────

@router.callback_query(F.data.startswith("qgroup:"))
async def group_play_quiz(callback: CallbackQuery, db_user: User):
    """Start a group quiz session."""
    parts = callback.data.split(":")
    quiz_id = int(parts[1])
    part_num = int(parts[2]) if len(parts) > 2 else 1
    
    quiz = await get_quiz_with_questions(quiz_id)
    if not quiz:
        await callback.answer("❌ Quiz topilmadi!", show_alert=True)
        return
    
    chat_id = callback.message.chat.id
    
    # Check if already running in this chat
    if chat_id in _group_sessions:
        await callback.answer("⚠️ Bu chatda quiz allaqachon ketayapti!", show_alert=True)
        return
    
    per_quiz = quiz.per_quiz
    start_idx = (part_num - 1) * per_quiz
    end_idx = min(part_num * per_quiz, len(quiz.questions))
    
    sorted_qs = sorted(quiz.questions, key=lambda q: q.order_num)
    chunk = sorted_qs[start_idx:end_idx]
    
    if not chunk:
        await callback.answer("❌ Savollar topilmadi!", show_alert=True)
        return
    
    play_questions = [
        {"question": qq.question, "options": list(qq.options), "correct_index": qq.correct_index}
        for qq in chunk
    ]
    
    if quiz.shuffle_mode == "all":
        random.shuffle(play_questions)
    
    # Create group session
    _group_sessions[chat_id] = {
        "quiz_id": quiz.id,
        "questions": play_questions,
        "index": 0,
        "total": len(play_questions),
        "time_per_q": quiz.time_per_q,
        "shuffle": quiz.shuffle_mode,
        "participants": {},
        "current_poll_id": None,
        "current_correct": None,
        "q_start_time": None,
        "answered_users": set(),
        "start_time": _time.time(),
    }
    
    await callback.answer()
    await callback.message.answer(
        f"📝 <b>Guruh quiz boshlandi!</b>\n"
        f"❓ Savollar: {len(play_questions)} ta | ⏱ Vaqt: {quiz.time_per_q} soniya\n"
        f"👥 Barcha qatnashchilar javob bering!\n\n"
        f"<i>Har bir savol {quiz.time_per_q} soniyadan keyin o'zgaradi.</i>",
        parse_mode="HTML"
    )
    
    await asyncio.sleep(2)
    await _send_group_next(callback.bot, chat_id)


# ══════════════════════════════════════════════════════════════════════════════
# POLL ANSWER HANDLER — works for both private and group quizzes
# ══════════════════════════════════════════════════════════════════════════════

@router.poll_answer()
async def handle_poll_answer(poll_answer: PollAnswer, state: FSMContext):
    """Handle user's answer to quiz poll — private & group."""
    user_id = poll_answer.user.id
    user_name = poll_answer.user.first_name or "User"
    
    # 1. Check group sessions
    for chat_id, session in list(_group_sessions.items()):
        if session.get("current_poll_id") == poll_answer.poll_id:
            if user_id in session.get("answered_users", set()):
                return  # Already answered
            
            session["answered_users"].add(user_id)
            
            # Initialize participant
            if user_id not in session["participants"]:
                session["participants"][user_id] = {
                    "name": user_name,
                    "correct": 0,
                    "total": 0,
                    "total_time": 0.0,
                }
            
            p = session["participants"][user_id]
            p["total"] += 1
            
            # Track time
            q_start = session.get("q_start_time", _time.time())
            answer_time = _time.time() - q_start
            p["total_time"] += answer_time
            
            # Check correctness
            correct_idx = session.get("current_correct", -1)
            if len(poll_answer.option_ids) > 0 and poll_answer.option_ids[0] == correct_idx:
                p["correct"] += 1
            
            return
    
    # 2. Check private (FSM) session
    current_state = await state.get_state()
    if current_state != QuizStates.playing.state:
        return
    
    data = await state.get_data()
    expected_poll_id = data.get("play_poll_id")
    
    if expected_poll_id and poll_answer.poll_id == expected_poll_id:
        correct_idx = data.get("play_current_correct", -1)
        play_correct = data.get("play_correct", 0)
        
        if len(poll_answer.option_ids) > 0 and poll_answer.option_ids[0] == correct_idx:
            play_correct += 1
            await state.update_data(play_correct=play_correct)
    
    # Signal the waiting event to move to next question
    event = _quiz_answer_events.get(user_id)
    if event:
        event.set()


# ══════════════════════════════════════════════════════════════════════════════
# MENING QUIZLARIM — list user's saved quizzes
# ══════════════════════════════════════════════════════════════════════════════

@router.message(F.text == "📂 Mening quizlarim")
async def my_quizzes(message: Message, state: FSMContext, db_user: User):
    quizzes = await get_user_quizzes(db_user.id)
    
    if not quizzes:
        await message.answer(
            "📂 <b>Sizda hali quiz yo'q.</b>\n\n"
            "📝 Quiz yaratish uchun fayl yuboring!",
            parse_mode="HTML",
            reply_markup=quiz_menu_kb()
        )
        return
    
    for quiz in quizzes[:10]:
        num_parts = math.ceil(quiz.total_questions / quiz.per_quiz) if quiz.per_quiz else 1
        date_str = quiz.created_at.strftime("%d.%m.%Y") if quiz.created_at else ""
        
        buttons = []
        if num_parts == 1:
            buttons.append([InlineKeyboardButton(text="▶️ Boshlash", callback_data=f"qplay:{quiz.id}:1")])
        else:
            for i in range(1, min(num_parts + 1, 6)):  # max 5 parts shown
                buttons.append([InlineKeyboardButton(
                    text=f"▶️ {quiz.name} {i}", callback_data=f"qplay:{quiz.id}:{i}"
                )])
        buttons.append([InlineKeyboardButton(text="🗑 O'chirish", callback_data=f"qdel:{quiz.id}")])
        
        await message.answer(
            f"📋 <b>{quiz.name}</b>\n"
            f"❓ Savollar: {quiz.total_questions} ta\n"
            f"📅 {date_str}",
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons)
        )


@router.callback_query(F.data.startswith("qdel:"))
async def delete_quiz_handler(callback: CallbackQuery, state: FSMContext, db_user: User):
    quiz_id = int(callback.data.split(":")[1])
    await delete_quiz(quiz_id)
    await callback.message.delete()
    await callback.answer("🗑 Quiz o'chirildi!")




# ══════════════════════════════════════════════════════════════════════════════
# STOP DURING QUIZ — show partial results
# ══════════════════════════════════════════════════════════════════════════════

@router.message(QuizStates.playing, Command("stop"))
@router.message(QuizStates.playing, F.text.in_({"🏠 Bosh menyu", "/start", "/cancel"}))
async def stop_playing_quiz(message: Message, state: FSMContext):
    """Stop quiz and show partial results."""
    data = await state.get_data()
    play_correct = data.get("play_correct", 0)
    play_index = data.get("play_index", 0)
    total = len(data.get("play_questions", []))
    start_time = data.get("play_start_time", _time.time())
    elapsed = _time.time() - start_time
    
    answered = play_index
    if answered > 0:
        pct = round(play_correct / answered * 100)
    else:
        pct = 0
    
    user_id = message.from_user.id
    _quiz_answer_events.pop(user_id, None)
    await state.clear()
    
    await message.answer(
        f"🛑 <b>Quiz to'xtatildi!</b>\n\n"
        f"┌─────────────────────\n"
        f"│ 📋 Jami savollar: <b>{total}</b> ta\n"
        f"│ ✏️ Javob berildi: <b>{answered}</b> ta\n"
        f"│ ✅ To'g'ri: <b>{play_correct}</b> ta\n"
        f"│ ❌ Noto'g'ri: <b>{answered - play_correct}</b> ta\n"
        f"│ 📊 Natija: <b>{pct}%</b>\n"
        f"│ ⏱ Vaqt: <b>{_format_time(elapsed)}</b>\n"
        f"└─────────────────────",
        parse_mode="HTML",
        reply_markup=quiz_menu_kb()
    )


# ══════════════════════════════════════════════════════════════════════════════
# WRONG INPUT
# ══════════════════════════════════════════════════════════════════════════════

@router.message(QuizStates.waiting_for_file, ~F.document)
async def quiz_wrong_input(message: Message, state: FSMContext, db_user: User):
    if message.text and ("Bosh menyu" in message.text or "menyu" in message.text.lower()):
        await state.clear()
        from handlers.start import build_welcome_text
        await message.answer(build_welcome_text(db_user), reply_markup=main_menu_kb(_lang(db_user)), parse_mode="HTML")
        return
    await message.answer(
        "📎 <b>Iltimos, fayl yuboring!</b>\n\n"
        "Qo'llab-quvvatlanadigan formatlar: DOCX, DOC, TXT, PDF",
        parse_mode="HTML"
    )
