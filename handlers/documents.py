from __future__ import annotations
import asyncio, json, os, io, logging
from aiogram import Router, F
from aiogram.types import Message, CallbackQuery, BufferedInputFile
from aiogram.exceptions import TelegramBadRequest
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

from database.models import User
from database.db import create_request, deduct_balance, mark_free_used
from keyboards.main_kb import main_menu_kb, back_to_menu_kb
from keyboards.documents_kb import page_count_kb, referat_summary_kb, doc_initial_settings_kb
from services.ai_service import generate_document_plan, generate_document_section
from services.docx_service import generate_docx, generate_docx_from_template
from services.template_loader import load_template_and_example
from utils.helpers import format_price, is_free_trial
from utils.i18n import t
from config import ADMIN_IDS

def _lang(db_user) -> str:
    return (db_user.language or "uz") if db_user else "uz"

logger = logging.getLogger(__name__)
router = Router()

class DocumentStates(StatesGroup):
    entering_topic    = State()
    entering_subject  = State()
    choosing_quality  = State()
    choosing_pages    = State()
    reviewing_summary = State()
    waiting_for_file      = State()   # Mustaqil ish: fayl kutilmoqda
    waiting_file_language = State()   # Mustaqil ish: fayl o'qildi, til kutilmoqda

# ─── Step 1: Start ──────────────────────────────────────────────────────────

from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from config import PRICING, ADMIN_IDS

@router.message(F.text.in_(["📚 Referat Yaratish", "📄 Mustaqil Ish Yaratish", "📚 Referat yaratish", "📄 Mustaqil ish yaratish", "📚 Referat", "📚 Essay"]))
async def start_referat_creation(message: Message, state: FSMContext, db_user: User):
    is_referat = "Referat" in message.text or "Referat" in message.text or "Essay" in message.text or "referat" in message.text.lower()
    service_label = "     📚 <b>Referat</b>" if is_referat else "    📄 <b>Mustaqil Ish</b>"
    price_std = PRICING["essay_std"] if is_referat else PRICING["mustaqil_std"]
    price_pro = PRICING["essay_pre"] if is_referat else PRICING["mustaqil_pre"]
    service_type = "referat" if is_referat else "mustaqil"

    is_admin = db_user.id in ADMIN_IDS
    free_trial = is_free_trial(db_user) and not is_admin
    balance = db_user.balance or 0
    if free_trial:
        balance_display = "🎁 Birinchi marta BEPUL!"
    elif is_admin:
        balance_display = "🛡️ Admin (tekin)"
    else:
        balance_display = format_price(balance)

    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Davom etish", callback_data=f"doc_confirm:{service_type}")],
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="doc_cancel")]
    ])

    await message.answer(
        f"{service_label}\n\n"
        f"────────────────\n"
        f"💰 <b>Narx:</b> 5 000 so'm — 20 000 so'm\n"
        f"   <i>(Sahifalar soniga qarab)</i>\n"
        f"────────────────\n",
        reply_markup=kb,
        parse_mode="HTML"
    )

@router.callback_query(F.data.startswith("doc_confirm:"))
async def confirm_document(callback: CallbackQuery, state: FSMContext, db_user: User):
    service_type = callback.data.split(":")[1]  # 'referat' or 'mustaqil'
    await state.update_data(
        service_type=service_type,
        is_referat=(service_type == "referat"),
        author=db_user.full_name or "Foydalanuvchi"
    )
    await state.set_state(DocumentStates.reviewing_summary)
    await callback.message.delete()
    await callback.message.answer(
        "⚙️ <b>Kerakli sozlamalaringizni to'ldiring:</b>\n\n"
        "<i>Mavzu, fan, sifat, sahifalar soni va boshqa ma'lumotlarni webapp orqali kiriting.</i>",
        reply_markup=doc_initial_settings_kb(
            db_user.balance or 0, db_user.full_name or "", doc_type=service_type, lang=_lang(db_user)
        ),
        parse_mode="HTML"
    )
    await callback.answer()

@router.callback_query(F.data == "doc_cancel")
async def cancel_document(callback: CallbackQuery, state: FSMContext):
    await state.clear()
    await callback.message.edit_text(t("cancelled", "uz"))
    await callback.answer()

@router.callback_query(F.data == "doc_topup")
async def topup_from_document(callback: CallbackQuery, state: FSMContext):
    await state.clear()
    await callback.message.edit_text(
        t("pres_topup_hint", "uz"),
        parse_mode="HTML"
    )
    await callback.answer()


# ─── File upload flow (Mustaqil ish only) ───────────────────────────────────

async def _read_file(message: Message) -> str:
    """Extract text from .txt / .docx / .pdf document."""
    doc  = message.document
    file = await message.bot.get_file(doc.file_id)
    buf  = io.BytesIO()
    await message.bot.download_file(file.file_path, buf)
    buf.seek(0)
    name = (doc.file_name or "").lower()
    if name.endswith(".txt"):
        return buf.read().decode("utf-8", errors="ignore")
    elif name.endswith(".docx"):
        from docx import Document as DocxDoc
        return "\n".join(p.text for p in DocxDoc(buf).paragraphs if p.text.strip())
    elif name.endswith(".pdf"):
        import pdfplumber
        parts = []
        with pdfplumber.open(buf) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n".join(parts)
    else:
        raise ValueError("Qo'llab-quvvatlanmaydigan format. .txt, .docx yoki .pdf yuboring.")


@router.callback_query(F.data.startswith("doc_file:"))
async def ask_for_file(callback: CallbackQuery, state: FSMContext):
    service_type = callback.data.split(":")[1]   # 'mustaqil'
    await state.update_data(service_type=service_type, is_referat=False)
    await state.set_state(DocumentStates.waiting_for_file)
    await callback.message.edit_text(
        t("doc_send_file", "uz"),
        parse_mode="HTML"
    )
    await callback.answer()


@router.message(F.document, DocumentStates.waiting_for_file)
async def receive_file_for_mustaqil(message: Message, state: FSMContext, db_user: User):
    proc = await message.answer(t("tfill_file_read", "uz"), parse_mode="HTML")
    try:
        source_text = await _read_file(message)
        if not source_text.strip():
            await proc.edit_text(t("doc_file_empty", "uz"))
            return
        source_text = source_text[:15000]

        # Detect complexity marker
        lower = source_text.lower()
        is_murakkab = any(kw in lower for kw in [
            "murakkab", "мураккаб", "сложный", "complex", "qiyin topshiriq"
        ])
        price    = 7000 if is_murakkab else 5000
        mode     = "murakkab" if is_murakkab else "normal"
        balance  = db_user.balance or 0
        is_admin = db_user.id in ADMIN_IDS
        free_trial = is_free_trial(db_user) and not is_admin
        if free_trial:
            price = 0

        if not is_admin and not free_trial and balance < price:
            await proc.edit_text(
                f"⚠️ <b>{'Murakkab topshiriq' if is_murakkab else 'Mustaqil ish'} aniqlandi!</b>\n"
                f"💰 Narxi: <b>{format_price(price)}</b>\n"
                f"💳 Balansingiz: <b>{format_price(balance)}</b>\n\n"
                "❌ Balans yetarli emas.",
                parse_mode="HTML"
            )
            await state.clear()
            return

        await state.update_data(
            source_text=source_text,
            quality="pro",
            mode=mode,
            price=price,
            num_pages=10,        # default, AI decides actual length from file
            topic="avtomatik",   # AI extracts from file
            author=db_user.full_name or "Foydalanuvchi",
        )
        await state.set_state(DocumentStates.waiting_file_language)
        await proc.delete()

        badge = "🔴 <b>Murakkab topshiriq!</b>" if is_murakkab else "✅ <b>Fayl o'qildi!</b>"
        await message.answer(
            f"{badge}\n"
            f"📝 Hajm: ~{len(source_text.split())} so'z\n"
            f"💰 Narxi: <b>{format_price(price)}</b>\n\n"
            + t("doc_choose_lang", "uz"),
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="🇺🇿 O'zbekcha",    callback_data="file_lang:uz")],
                [InlineKeyboardButton(text="🇷🇺 Русский",      callback_data="file_lang:ru")],
                [InlineKeyboardButton(text="🇬🇧 English",       callback_data="file_lang:en")],
                [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="file_lang:cancel")],
            ]),
            parse_mode="HTML"
        )
    except ValueError as e:
        await proc.edit_text(f"❌ {e}")
    except Exception as e:
        await proc.edit_text(f"❌ Xatolik: {e}")


@router.callback_query(F.data.startswith("file_lang:"), DocumentStates.waiting_file_language)
async def file_language_chosen(callback: CallbackQuery, state: FSMContext, db_user: User):
    lang = callback.data.split(":")[1]
    if lang == "cancel":
        await state.clear()
        await callback.message.edit_text(t("cancelled", "uz"))
        await callback.answer()
        return

    await state.update_data(language=lang)
    data       = await state.get_data()
    source_text = data["source_text"]
    mode        = data["mode"]
    price       = data["price"]
    is_admin    = db_user.id in ADMIN_IDS

    lang_label = {"uz": "O'zbekcha 🇺🇿", "ru": "Rus tili 🇷🇺", "en": "English 🇬🇧"}.get(lang, lang)
    lang_map   = {"uz": "O'zbek tilida (Lotin alifbosi)", "ru": "На русском языке", "en": "In English"}
    lang_instruction = lang_map.get(lang, "O'zbek tilida")

    await callback.message.edit_text(
        f"{'🔴 Murakkab' if mode == 'murakkab' else '🧠 AI'} rejim | {lang_label}\n"
        f"🔄 Tahlil boshlandi...",
        parse_mode="HTML"
    )
    await callback.answer()

    wait_msg = await callback.message.bot.send_message(
        chat_id=callback.message.chat.id,
        text=t("doc_analyzing", lang),
        parse_mode="HTML"
    )

    try:
        from services.ai_service import _call_ai

        async def call_openai_with_retry(messages, temp, max_tok, retries=4):
            return await _call_ai(messages, max_tokens=max_tok, temperature=temp, retries=retries)

        # ── STEP 1: Extract sections from file ────────────────────────────────
        sections_raw = await call_openai_with_retry(
            messages=[
                {"role": "system", "content":
                    "You are an academic document analyzer. "
                    "Read the file content and extract the exact list of sections that need to be written based ONLY on the user's instructions. "
                    "Output ONLY section names, one per line. No numbers, no extra text. If no explicit structure is given, derive a logical detailed structure for a comprehensive assignment."
                },
                {"role": "user", "content":
                    f"FILE CONTENT:\n{source_text[:8000]}\n\n"
                    f"Task: Extract or design the section plan. Language for section names: {lang_instruction}."}
            ],
            temp=0.3,
            max_tok=800
        )
        sections = [s.strip().lstrip("0123456789.-) ") for s in sections_raw.split("\n") if s.strip()]
        if not sections:
            sections = ["Kirish", "Asosiy qism", "Xulosa"]

        total_sec = len(sections)

        await wait_msg.edit_text(
            f"🧠 <b>AI tahlil qildi!</b>\n\n"
            f"📋 <b>{total_sec} ta bo'lim aniqlandi</b>\n"
            + "\n".join(f"  {i+1}. {s}" for i, s in enumerate(sections))
            + "\n\n<i>Yozish boshlandi...</i>",
            parse_mode="HTML"
        )
        await asyncio.sleep(1.5)

        # ── STEP 2: Section-by-section generation ────────────────────────────
        murakkab_rule = (
            "CRITICAL MODE: Follow file instructions EXACTLY and completely. "
            "No templates, no extra rules, no simplifications. "
            "Do exactly and fully what the file says.\n\n"
        ) if mode == "murakkab" else ""

        system_prompt = (
            "You are a world-class academic expert executing a complex assignment. "
            "You MUST fulfill the user's file instructions with absolute precision and completely unrestricted depth. "
            "If the file asks for 4000, 6000 words, or extensive details, YOU MUST DO IT. There are no length limits. "
            "Every section must be comprehensive, analytical, and highly detailed.\n\n"
            f"{murakkab_rule}"
            f"DOCUMENT LANGUAGE: {lang_instruction}\n\n"
            f"MANDATORY RULES:\n"
            "- Never shorten or summarize unless requested.\n"
            "- Do NOT use markdown symbols (**, #, ```).\n"
            "- Write ONLY in full academic paragraphs.\n"
            "- Do NOT write section titles in the body — they are headings.\n\n"
            f"FULL FILE CONTENT (your primary reference for ALL sections):\n"
            f"{'='*60}\n{source_text[:10000]}\n{'='*60}"
        )

        full_parts = []

        for i, section_title in enumerate(sections, 1):
            done = "🟩" * (i - 1)
            cur  = "🟨"
            left = "⬜" * (total_sec - i)
            bar  = done + cur + left

            # Phase 1: Planning
            try:
                await wait_msg.edit_text(
                    f"🧠 <b>AI chuqur tahlil qilmoqda...</b>\n\n"
                    f"📄 <b>Bo'lim {i}/{total_sec}</b>\n"
                    f"{bar}\n\n"
                    f"🔍 <b>1-bosqich — Reja:</b> <i>{section_title}</i>\n\n"
                    f"<i>AI avval fikrlaydi, keyin yozadi — sifat uchun sabr qiling...</i>",
                    parse_mode="HTML"
                )
            except:
                pass

            planning_prompt = (
                f"You are preparing to write the '{section_title}' section.\n\n"
                "Perform a DEEP pre-writing analysis based on the file content:\n"
                "1. What are the 5 most critical points for this section?\n"
                "2. What specific facts, statistics, or examples from the file apply here?\n"
                "3. What additional academic knowledge supports this section?\n"
                "4. What is the most logical paragraph order for maximum impact?\n"
                "5. What counterarguments exist and how to address them?\n\n"
                "Write your analysis plan (250–400 words)."
            )
            try:
                section_plan = await call_openai_with_retry(
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user",   "content": planning_prompt}
                    ],
                    temp=0.5,
                    max_tok=2000
                )
            except:
                section_plan = ""

            # Phase 2: Writing
            try:
                await wait_msg.edit_text(
                    f"✍️ <b>AI yozmoqda...</b>\n\n"
                    f"📄 <b>Bo'lim {i}/{total_sec}</b>\n"
                    f"{bar}\n\n"
                    f"📝 <b>2-bosqich — Yozish:</b> <i>{section_title}</i>\n\n"
                    f"<i>Reja tayyor — to'liq matn yozilmoqda...</i>",
                    parse_mode="HTML"
                )
            except:
                pass
            await asyncio.sleep(0.5)

            write_prompt = (
                f"Write the '{section_title}' section of the academic document.\n\n"
                f"MANDATORY:\n"
                f"- Obey ALL instructions from the file content.\n"
                f"- Write in {lang_instruction}.\n"
                f"- DO NOT artificially limit the length. If the assignment requires huge detail, provide it.\n"
                f"- Do NOT write the section title — only the content.\n"
                f"- No markdown. Only full academic paragraphs.\n"
                + (f"\nYOUR PRE-WRITING ANALYSIS (use as blueprint):\n{section_plan}" if section_plan else "")
            )

            section_text = await call_openai_with_retry(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": write_prompt}
                ],
                temp=0.8,
                max_tok=15000
            )
            full_parts.append(f"# {section_title}\n\n{section_text}")

        # Done
        try:
            await wait_msg.edit_text(
                f"✅ <b>Barcha {total_sec} ta bo'lim tayyor!</b>\n\n"
                f"{'🟩' * total_sec}\n\n"
                f"⏳ <b>DOCX fayl yaratilmoqda...</b>",
                parse_mode="HTML"
            )
        except:
            pass

        content = "\n\n".join(full_parts)
        topic_label = "Fayldan yaratilgan mustaqil ish"

        docx_bytes = await asyncio.get_event_loop().run_in_executor(
            None, generate_docx, "report", topic_label,
            content, db_user.full_name or "Foydalanuvchi", ""
        )

        if free_trial:
            await mark_free_used(db_user.id)
        elif not is_admin:
            await deduct_balance(db_user.id, price)

        await create_request(db_user.id, "essay", topic=topic_label)
        await wait_msg.delete()
        await callback.message.bot.send_document(
            chat_id=callback.message.chat.id,
            document=BufferedInputFile(docx_bytes, filename="mustaqil_ish.docx"),
            caption=(
                f"✅ <b>Mustaqil ish tayyor!</b>\n"
                f"📎 Fayl asosida yaratildi ({total_sec} bo'lim)\n"
                f"🌐 Til: {lang_label}\n"
                f"💰 {'Hisobdan yechildi: ' + format_price(price) if not is_admin else '🛡️ Admin — tekin'}"
            ),
            reply_markup=main_menu_kb(),
            parse_mode="HTML"
        )
        await state.clear()

    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"file_language_chosen error: {e}")
        try:
            await wait_msg.edit_text(f"❌ Xatolik yuz berdi: {e}")
        except:
            await callback.message.bot.send_message(callback.message.chat.id, f"❌ Xatolik: {e}")
        await state.clear()


@router.message(DocumentStates.entering_topic)
async def enter_topic(message: Message, state: FSMContext):
    await state.update_data(topic=message.text.strip())
    await state.set_state(DocumentStates.entering_subject)
    await message.answer(
        t("doc_enter_subject", "uz"),
        reply_markup=back_to_menu_kb(),
        parse_mode="HTML"
    )

@router.message(DocumentStates.entering_subject)
async def enter_subject(message: Message, state: FSMContext):
    await state.update_data(subject=message.text.strip())
    await state.set_state(DocumentStates.choosing_quality)
    from keyboards.documents_kb import referat_quality_kb
    await message.answer(
        t("doc_choose_quality", "uz"),
        reply_markup=referat_quality_kb(),
        parse_mode="HTML"
    )

@router.callback_query(F.data.startswith("ref_quality:"), DocumentStates.choosing_quality)
async def choose_quality(callback: CallbackQuery, state: FSMContext):
    quality = callback.data.split(":")[1]
    is_pro = (quality == "pro")
    await state.update_data(quality=quality, is_pro=is_pro)
    await state.set_state(DocumentStates.choosing_pages)
    from keyboards.documents_kb import page_count_kb
    await callback.message.edit_text(
        f"📑 <b>{'💎 Pro' if is_pro else '✨ Standart'} | Hajmni tanlang:</b>",
        reply_markup=page_count_kb(is_pro=is_pro),
        parse_mode="HTML"
    )
    await callback.answer()

@router.callback_query(F.data.startswith("pages:"), DocumentStates.choosing_pages)
async def choose_pages(callback: CallbackQuery, state: FSMContext, db_user: User):
    _, pages, price = callback.data.split(":")
    await state.update_data(num_pages=int(pages), price=int(price))
    
    # Auto-generate a plan first
    data = await state.get_data()
    topic = data.get("topic")
    
    wait_msg = await callback.message.edit_text(
        t("doc_preparing_plan", "uz"),
        parse_mode="HTML"
    )
    
    try:
        # For Pro version, we explicitly ask for a much longer plan
        is_pro = data.get("is_pro", False)
        plan_prompt = "7-10 ta juda batafsil bo'lim" if is_pro else "5-6 ta bo'lim"
        
        real_service = data.get("service_type", "essay")
        plan_data = await generate_document_plan(
            service_type=real_service, 
            topic=topic, 
            language=data.get("language", "uz"),
            detail_level="pro" if is_pro else "standard"
        )
        plan = plan_data.get("plan", "")
        sections = [s.strip() for s in plan.split('\n') if s.strip()]
        
        await state.update_data(manual_plan=plan, num_chapters=len(sections))
    except Exception as e:
        logger.error(f"Document plan generation error: {e}")
    
    await wait_msg.delete()
    await state.set_state(DocumentStates.reviewing_summary)
    await show_referat_summary(callback.message, state, db_user)
    await callback.answer()

@router.message(F.web_app_data, DocumentStates.reviewing_summary)
async def handle_referat_webapp(message: Message, state: FSMContext, db_user: User):
    try:
        try:
            await message.delete()
        except TelegramBadRequest:
            pass
        data = json.loads(message.web_app_data.data)
        action = data.get("action", "")
        
        if action == "plan_update":
            await state.update_data(manual_plan=data.get("manual_plan"), num_chapters=data.get("num_chapters"))
            await show_referat_summary(message, state, db_user)
        elif action == "content_update":
            await state.update_data(extra_info=data.get("extra_info"), tone=data.get("tone"))
            await show_referat_summary(message, state, db_user)
        elif action == "full_settings":
            # Full settings from webapp — topic, subject, quality, pages, price, etc.
            quality = data.get("quality", "standard")
            ai_images = data.get("ai_images", 3)
            await state.update_data(
                topic=data.get("topic"),
                subject=data.get("subject"),
                author=data.get("author", db_user.full_name or "Foydalanuvchi"),
                language=data.get("language", "uz"),
                quality=quality,
                is_pro=(quality == "pro"),
                num_pages=data.get("num_pages", 15),
                price=data.get("price", 3000),
                base_price=data.get("price", 3000),
                ai_images_count=ai_images,
                fakultet=data.get("fakultet", ""),
                specialty=data.get("specialty", ""),
                reviewer=data.get("reviewer", ""),
                university=data.get("university", ""),
            )
            # Auto-generate plan
            topic = data.get("topic")
            wait_msg = await message.answer(
                t("doc_preparing_plan", "uz"),
                parse_mode="HTML"
            )
            try:
                real_service = data.get("service_type", "essay")
                plan_data = await generate_document_plan(
                    service_type=real_service,
                    topic=topic,
                    language=data.get("language", "uz"),
                    detail_level="pro" if quality == "pro" else "standard"
                )
                plan = plan_data.get("plan", "")
                sections = [s.strip() for s in plan.split('\n') if s.strip()]
                await state.update_data(manual_plan=plan, num_chapters=len(sections))
            except Exception as e:
                logger.error(f"Plan generation in webapp error: {e}")
            await wait_msg.delete()
            await show_referat_summary(message, state, db_user)
        else:
            # Legacy settings update (author, language only)
            if "author" in data:
                await state.update_data(author=data.get("author"), language=data.get("language"))
            await show_referat_summary(message, state, db_user)
    except json.JSONDecodeError as e:
        logger.error(f"Referat WebApp JSON decode error: {e}")
        await message.answer("⚠️ Ma'lumotlarni o'qishda xatolik. Qaytadan urinib ko'ring.")
    except Exception as e:
        logger.error(f"Referat WebApp error: {e}")
        await message.answer("⚠️ Kutilmagan xatolik yuz berdi.")

async def show_referat_summary(message: Message, state: FSMContext, db_user: User):
    import html as _html
    data = await state.get_data()
    topic = _html.escape(data.get("topic") or "")
    subject = _html.escape(data.get("subject") or "")
    pages = data.get("num_pages")
    price = data.get("price")
    author = _html.escape(data.get("author", "Foydalanuvchi"))
    quality_val = data.get("quality", "standard")
    quality_label = "💎 Pro" if quality_val == "pro" else "✨ Standart"
    
    text = (
        f"┌─ 📝 <b>HUJJAT</b> │ {quality_label} ─┐\n\n"
        f"  📚  <b>Mavzu:</b> {topic}\n"
        f"  🎓  <b>Fan:</b> {subject}\n"
        f"  👤  <b>Muallif:</b> {author}\n"
        f"  📄  <b>Sahifalar:</b> {pages} bet\n"
        f"  💰  <b>Narxi:</b> {format_price(price)}\n\n"
        f"└──────────────────────┘\n\n"
    )
    is_admin = db_user.id in ADMIN_IDS
    from utils.helpers import is_free_trial
    free_trial = is_free_trial(db_user) and not is_admin
    balance = db_user.balance or 0
    
    if is_admin or free_trial or balance >= price:
        text += f"<i>Reja yoki ma'lumotlarni o'zgartirish uchun tugmalardan foydalaning.</i>"
    else:
        needed = price - balance
        from handlers.payment import get_cards_text
        text += (
            f"⚠️ <b>Sizning hisobingizda mablag' yetarli emas!</b>\n"
            f"🔴 Yetishmayapti: <b>{format_price(needed)}</b>\n\n"
            f"{get_cards_text()}\n"
            f"🧾 <i>Chekni to'g'ridan-to'g'ri shu yerga yuborishingiz mumkin.</i>"
        )
    
    summary_id = data.get("summary_msg_id")
    if summary_id:
        try:
            await message.bot.delete_message(chat_id=message.chat.id, message_id=summary_id)
        except TelegramBadRequest:
            pass

    doc_type = data.get("service_type", "referat")
    msg = await message.answer(text, reply_markup=referat_summary_kb(
        balance=db_user.balance or 0, price=price, name=db_user.full_name or "",
        topic=topic, doc_type=doc_type, quality=quality_val, lang=_lang(db_user),
        is_admin=(db_user.id in ADMIN_IDS)
    ), parse_mode="HTML")
    await state.update_data(summary_msg_id=msg.message_id)

@router.message(DocumentStates.reviewing_summary, F.photo)
async def receive_doc_receipt(message: Message, state: FSMContext, bot: Bot, db_user: User):
    photo = message.photo[-1]
    file_id = photo.file_id
    
    data = await state.get_data()
    
    # We need to know the price
    doc_type = data.get("service_type", "referat")
    quality = data.get("quality", "standard")
    pages = int(data.get("num_pages", 15))
    
    # Calculate price based on page tiers
    if pages <= 10:
        price = 5000
    elif pages <= 20:
        price = 10000
    elif pages <= 25:
        price = 15000
    else:
        price = 20000
        
    price += data.get("ai_images_cost", 0)
    
    wait_msg = await message.answer("⏳ Chek tekshirilmoqda. Iltimos biroz kuting...")
    
    try:
        import base64
        from io import BytesIO
        file_info = await bot.get_file(file_id)
        file_bytes = BytesIO()
        await bot.download_file(file_info.file_path, file_bytes)
        base64_img = base64.b64encode(file_bytes.getvalue()).decode("utf-8")
        
        from services.ai_service import verify_receipt_with_ai
        result = await verify_receipt_with_ai(base64_img, price)
        
        await wait_msg.delete()
        
        if result.get("verified") and result.get("amount", 0) >= price:
            from database.db import add_balance
            db_user = await add_balance(db_user.id, price)
            await message.answer(f"✅ To'lov muvaffaqiyatli tasdiqlandi! Balansingizga {format_price(price)} qo'shildi.\nEndi <b>✅ Yaratish</b> tugmasini bosishingiz mumkin.", parse_mode="HTML")
            
            # Re-render the summary message to show Yaratish button
            from keyboards.documents_kb import referat_summary_kb
            from utils.i18n import t
            
            text = (
                f"📝 <b>HUJJAT | ✨ {quality.capitalize()}</b>\n\n"
                f"📚 <b>Mavzu:</b> {data.get('topic')}\n"
                f"🎓 <b>Fan:</b> {data.get('subject')}\n"
                f"👤 <b>Muallif:</b> {data.get('author')}\n"
                f"📄 <b>Sahifalar:</b> {pages} bet\n"
                f"💰 <b>Narxi:</b> {format_price(price)}\n\n"
                f"──────────\n\n"
                f"<i>Reja yoki ma'lumotlarni o'zgartirish uchun tugmalardan foydalaning.</i>"
            )
            msg = await message.answer(text, reply_markup=referat_summary_kb(
                balance=db_user.balance or 0, price=price, name=db_user.full_name or "",
                topic=data.get('topic'), doc_type=doc_type, quality=quality, lang=db_user.language or "uz",
                is_admin=(db_user.id in ADMIN_IDS)
            ), parse_mode="HTML")
            await state.update_data(summary_msg_id=msg.message_id)
            return
        else:
            reason = result.get("reason", "Noma'lum xatolik")
            await message.answer(f"⚠️ Chekni avtomatik tasdiqlab bo'lmadi.\nSabab: <i>{reason}</i>\n\nChek adminga qo'lda tasdiqlash uchun yuborilmoqda...", parse_mode="HTML")
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"Error checking receipt: {e}")
        try:
            await wait_msg.delete()
        except:
            pass
        await message.answer("⚠️ Chekni avtomatik tekshirishda xatolik yuz berdi. Adminga yuborilmoqda...")
        
    from database.db import create_payment
    payment = await create_payment(user_id=db_user.id, amount=price, package=f"doc_{doc_type}", screenshot_file_id=file_id)
    
    from keyboards.admin_kb import payment_action_kb
    admin_text = (
        f"💰 <b>Yangi to'lov so'rovi! (Hujjat)</b>\n\n"
        f"👤 Foydalanuvchi: {db_user.full_name}\n"
        f"🆔 ID: <code>{db_user.id}</code>\n"
        f"💵 Summa: {format_price(price)}\n"
        f"🧾 ID: #{payment.id}"
    )
    
    from config import ADMIN_IDS
    for admin_id in ADMIN_IDS:
        try:
            await bot.send_photo(chat_id=admin_id, photo=file_id, caption=admin_text, reply_markup=payment_action_kb(payment.id, db_user.id), parse_mode="HTML")
        except Exception:
            pass

@router.message(F.text.contains("AI Rasm"), DocumentStates.reviewing_summary)
async def choose_ai_images_referat(message: Message, state: FSMContext, db_user: User):
    """Let user choose how many AI images to add (0-10, 1000 so'm each)."""
    buttons = []
    buttons.append([InlineKeyboardButton(text="❌ Rasmsiz (0)", callback_data="ref_img:0")])
    for i in range(1, 11, 2):
        row = []
        row.append(InlineKeyboardButton(text=f"🖼 {i} ta (+{i*1000} so'm)", callback_data=f"ref_img:{i}"))
        if i + 1 <= 10:
            row.append(InlineKeyboardButton(text=f"🖼 {i+1} ta (+{(i+1)*1000} so'm)", callback_data=f"ref_img:{i+1}"))
        buttons.append(row)
    kb = InlineKeyboardMarkup(inline_keyboard=buttons)
    data = await state.get_data()
    current = data.get("ai_images_count", 0)
    await message.answer(
        f"� <b>AI rasm qo'shish</b>\n\n"
        f"Hozirgi: {current} ta rasm\n"
        f"Har bir rasm: 1 000 so'm\n\n"
        f"Nechta rasm kerak?",
        reply_markup=kb, parse_mode="HTML"
    )

@router.callback_query(F.data.startswith("ref_img:"))
async def set_ai_images_referat(callback: CallbackQuery, state: FSMContext, db_user: User):
    count = int(callback.data.split(":")[1])
    data = await state.get_data()
    base_price = data.get("base_price", data.get("price", 3000))
    new_price = base_price + (count * 1000)
    await state.update_data(ai_images_count=count, price=new_price, base_price=base_price)
    await callback.message.delete()
    await callback.answer(f"✅ {count} ta AI rasm tanlandi" if count > 0 else "✅ Rasmsiz")
    await show_referat_summary(callback.message, state, db_user)

@router.message(F.text.in_({"✅ Yaratish", "🚀 Yaratish", " Yaratish"}), DocumentStates.reviewing_summary)
async def generate_referat(message: Message, state: FSMContext, db_user: User):
    data = await state.get_data()
    price = data.get("price", 3000)
    topic = data.get("topic")
    pages = data.get("num_pages")
    plan = data.get("manual_plan", "")
    
    lang = _lang(db_user)
    is_admin = db_user.id in ADMIN_IDS
    free_trial = is_free_trial(db_user) and not is_admin
    if free_trial:
        price = 0
    if not is_admin and not free_trial and (db_user.balance or 0) < price:
        await message.answer(t("pres_balance_ok", lang).replace("✅", "⚠️").replace("yetarli", "yetarli emas"), reply_markup=main_menu_kb())
        return
    # Cancellation support
    cancel_flag = {"cancelled": False}
    await state.update_data(cancel_flag=cancel_flag)
    
    cancel_kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="doc_cancel_gen")]
    ])

    wait_msg = await message.answer(
        t("doc_generating", lang, topic=topic, pages=pages),
        parse_mode="HTML",
        reply_markup=cancel_kb
    )

    def is_cancelled():
        return cancel_flag.get("cancelled", False)

    try:
        # Dynamic word count calculation to hit the page target
        sections = [s.strip() for s in plan.split('\n') if s.strip()] if plan else [topic]
        num_sections = len(sections)
        # 1 page is ~250-280 words with our 14pt 1.5 spacing format
        # Use 300 to overshoot slightly and guarantee page count
        total_target_words = pages * 300 
        words_per_section = int(total_target_words / num_sections)
        
        # Load structure template + quality example
        doc_type = data.get("service_type", "referat")
        template_context = load_template_and_example(doc_type)

        source_text = data.get("source_text", "")
        mode        = data.get("mode", "normal")

        full_content = []

        if mode == "murakkab" and source_text:
            # ── MURAKKAB: execute file instructions exactly, no templates ──────
            import logging
            logging.getLogger(__name__).info(f"MURAKKAB mode: {topic}")
            try:
                await wait_msg.edit_text(
                    "🔴 <b>Murakkab topshiriq bajarilmoqda...</b>\n\n"
                    "📄 AI fayldagi barcha ko'rsatmalarni so'zsiz bajarmoqda.\n"
                    "<i>Biroz sabr qiling...</i>",
                    parse_mode="HTML"
                )
            except:
                pass
            result = await generate_document_section(
                topic=topic,
                section_title="To'liq mustaqil ish",
                extra_details=(
                    f"FOYDALANUVCHI FAYLI — barcha ko'rsatmalar:\n{source_text}\n\n"
                    "QOIDA: Yuqoridagi faylda yozilgan BARCHA ko'rsatmalarni "
                    "SO'ZSIZ bajaring. Hech qanday shablon, qo'shimcha qoida yoki "
                    "cheklov ishlatmang. Faqat faylda yozilganidek, aynan o'sha "
                    f"format va tuzilmada bajaring. Hajm: ~{pages * 280} so'z."
                ),
                language=data.get("language", "uz"),
                quality="pro",
                service_type=doc_type
            )
            full_content = [result]

        else:
            # ── NORMAL: template-guided section generation ────────────────────
            source_block = (
                f"\n\n--- YUBORILGAN FAYL MATNI (asosiy manba) ---\n"
                f"{source_text[:5000]}\n"
                "--- FAYL MATNI TUGADI ---\n\n"
                "MUHIM: Faylda berilgan ma'lumotlar asosida yozing.\n"
            ) if source_text else ""

            for i, sec in enumerate(sections, 1):
                if is_cancelled():
                    await state.clear()
                    return
                
                if "ADABIYOT" in sec.upper():
                    continue
                    
                filled = int(((i - 1) / len(sections)) * 10)
                bar = "🟩" * filled + "⬜" * (10 - filled)
                try:
                    await wait_msg.edit_text(
                        f"🤖 <b>AI hujjat yozmoqda...</b>\n\n"
                        f"📊 <b>{i-1}/{len(sections)}</b> bo'lim tayyor\n"
                        f"{bar}\n\n"
                        f"⏭ <b>Hozir yozilmoqda:</b> <i>{sec}</i>\n\n"
                        f"⏳ <i>4-5 daqiqa kutib turing — sifatli natija uchun!</i>",
                        parse_mode="HTML",
                        reply_markup=cancel_kb
                    )
                except:
                    pass

                text = await generate_document_section(
                    topic=topic,
                    section_title=sec,
                    extra_details=(
                        f"{template_context}\n" if template_context else ""
                    ) + source_block + f"Ushbu bo'lim aynan {words_per_section} ta so'zdan iborat bo'lsin. " + data.get("extra_info", ""),
                    language=data.get("language", "uz"),
                    quality=data.get("quality", "standard"),
                    service_type=doc_type
                )
                full_content.append(f"# {sec}\n\n{text}")


        # Add bibliography
        try:
            bib_title = "FOYDALANILGAN ADABIYOTLAR RO'YXATI"
            try:
                await wait_msg.edit_text(f"⏳ <b>Hujjat yaratilmoqda...</b>\n📚 <i>Adabiyotlar ro'yxati shakllantirilmoqda...</i>")
            except TelegramBadRequest:
                pass
            
            bib_count = "10-15" if data.get("quality") == "pro" else "5-8"
            
            bib_text = await generate_document_section(
                topic=topic, 
                section_title=bib_title, 
                extra_details=(
                    f"Mavzu bo'yicha kamida {bib_count} ta REAL ilmiy manbalarni yoz.\n"
                    "TARKIBI:\n"
                    "- 40% O'zbek mualliflari (O'zbekiston nashriyotlari: Fan, Iqtisod-Moliya, Sharq, TDIU). "
                    "Masalan: Karimov A.K., Abdullayev B.M., Xo'jayev N.R. kabi.\n"
                    "- 60% Xorijiy mualliflar (Pearson, McGraw-Hill, Cambridge, Springer). "
                    "Masalan: Samuelson P., Mankiw G., Porter M. kabi.\n"
                    "FORMAT: [1] Familiya I.O. Kitob nomi. — Shahar: Nashriyot, yil. — bet soni.\n"
                    "Faqat ro'yxat ber, boshqa hech narsa qo'shma. Izoh yoki sarlavha yozma."
                ),
                language=data.get("language", "uz"),
                quality="standard",
                service_type="article"
            )
            full_content.append(f"# {bib_title}\n\n{bib_text}")
        except Exception as bib_err:
            logger.error(f"Bibliography generation error: {bib_err}")
        
        content = "\n\n".join(full_content)
        doc_type = data.get("service_type", "referat")
        
        if doc_type in ["referat", "mustaqil"]:
            # Use ready-made template for referat and mustaqil ish
            docx_bytes = await asyncio.get_event_loop().run_in_executor(
                None,
                generate_docx_from_template,
                topic,
                content,
                data.get("author", "Foydalanuvchi"),
                plan,
                data.get("subject", ""),
                data.get("reviewer", ""),
                data.get("university", ""),
                data.get("fakultet", ""),
                data.get("specialty", ""),
                doc_type,
                pages
            )
        else:
            # Fallback for other document types
            docx_bytes = await asyncio.get_event_loop().run_in_executor(
                None, 
                generate_docx, 
                "report", 
                topic, 
                content,
                data.get("author", "Foydalanuvchi"),
                plan,
                {"num_pages": pages}
            )
        
        # Generate AI images and insert into DOCX
        ai_images_count = data.get("ai_images_count", 0)
        if ai_images_count > 0:
            try:
                try:
                    await wait_msg.edit_text("🧠 <b>AI rasmlar uchun kalit so'zlar aniqlanmoqda...</b>", parse_mode="HTML")
                except TelegramBadRequest:
                    pass

                from services.image_search_service import generate_image_gemini
                from services.ai_service import _call_ai
                from docx import Document as DocxDocument
                from docx.shared import Inches as DocxInches, Pt as DocxPt, RGBColor as DocxRGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io as _io

                # Step 1: AI generates smart keywords — only for CONTENT sections
                # Filter out Reja, Kirish, Xulosa, Adabiyotlar from plan
                skip_words = ["reja", "kirish", "xulosa", "adabiyot", "foydalanilgan", "mundarija",
                              "план", "введение", "заключение", "литератур", "содержание",
                              "plan", "introduction", "conclusion", "reference", "bibliography"]
                content_sections = []
                if plan:
                    for s in plan.split('\n'):
                        s = s.strip()
                        if s and not any(sw in s.lower() for sw in skip_words):
                            content_sections.append(s)
                sections_text = "\n".join(content_sections) if content_sections else topic

                kw_prompt = (
                    f"Mavzu: {topic}\n"
                    f"Hujjat asosiy bo'limlari:\n{sections_text}\n\n"
                    f"VAZIFA: Ushbu akademik hujjat uchun {ai_images_count} ta FOTOSURAT kerak.\n"
                    f"Har bir fotosurat uchun INGLIZ tilida aniq kalit so'z yoz.\n\n"
                    f"FAQAT JSON array qaytar:\n"
                    f'[{{"section": "Bo\'lim nomi", "keyword": "realistic photograph of military first aid bandaging equipment on table"}}]\n\n'
                    f"MUHIM QOIDALAR:\n"
                    f"- Kalit so'z 8-15 so'zdan iborat, FAQAT fotosurat tasvirlash uchun\n"
                    f"- 'photograph of', 'realistic photo of' so'zlari bilan boshlang\n"
                    f"- Rasmlar FAQAT mavzuga mos bo'lsin\n"
                    f"- Erkaklar ko'rsatilishi mumkin, ayollar YO'Q\n"
                    f"- Professional, harbiy, jiddiy uslub\n"
                    f"- Har bir rasm BOSHQASIDAN FARQ QILSIN\n"
                    f"- DIQQAT: Kirish, Reja, Xulosa, Adabiyotlar uchun rasm QOSHMA!\n"
                    f"- Faqat ASOSIY mazmunli bo'limlar uchun rasm bo'lsin\n"
                    f"- Aniq {ai_images_count} ta element bo'lsin"
                )
                try:
                    import json as _json
                    kw_raw = await _call_ai(
                        [{"role": "user", "content": kw_prompt}],
                        max_tokens=1500, temperature=0.5, json_mode=True
                    )
                    kw_clean = kw_raw.replace("```json", "").replace("```", "").strip()
                    kw_list = _json.loads(kw_clean)
                    if not isinstance(kw_list, list):
                        kw_list = []
                except Exception:
                    kw_list = []

                # Fallback keywords if AI fails
                if len(kw_list) < ai_images_count:
                    for i in range(len(kw_list), ai_images_count):
                        kw_list.append({
                            "section": f"Bo'lim {i+1}",
                            "keyword": f"realistic photograph of professional {topic} equipment and tools"
                        })

                # Step 2: Generate images with Gemini and insert into CONTENT area only
                doc = DocxDocument(_io.BytesIO(docx_bytes))

                # Find content paragraphs — skip Reja, Kirish, Xulosa, Adabiyotlar
                all_paragraphs = list(doc.paragraphs)
                content_paragraphs = []
                in_content = False
                for p in all_paragraphs:
                    txt = p.text.strip().lower()
                    if not txt:
                        continue
                    # Detect start of main content (after Kirish)
                    is_skip = any(sw in txt for sw in skip_words)
                    if txt.startswith("1.") or txt.startswith("1 ") or txt.startswith("ii"):
                        in_content = True
                    # Detect end of content (Xulosa, Adabiyotlar)
                    if in_content and any(sw in txt for sw in ["xulosa", "заключение", "conclusion",
                                                                "adabiyot", "foydalanilgan", "литератур",
                                                                "reference", "bibliography"]):
                        in_content = False
                        continue
                    if in_content and not is_skip and len(txt) > 20:
                        content_paragraphs.append(p)

                # Fallback: if detection failed, use middle 60% of paragraphs
                if len(content_paragraphs) < 3:
                    total_p = [p for p in all_paragraphs if p.text.strip() and len(p.text.strip()) > 20]
                    start = max(1, len(total_p) // 5)       # skip first 20%
                    end = max(start + 3, len(total_p) * 4 // 5)  # skip last 20%
                    content_paragraphs = total_p[start:end]

                step = max(1, len(content_paragraphs) // (ai_images_count + 1))
                images_done = 0

                for img_i, kw_item in enumerate(kw_list[:ai_images_count]):
                    keyword = kw_item.get("keyword", f"realistic photograph of {topic}")
                    sec_name = kw_item.get("section", f"Bo'lim {img_i+1}")

                    bar_done = "🟩" * (img_i + 1)
                    bar_left = "⬜" * (ai_images_count - img_i - 1)
                    try:
                        await wait_msg.edit_text(
                            f"🎨 <b>Rasm {img_i+1}/{ai_images_count}</b>\n"
                            f"{bar_done}{bar_left}\n\n"
                            f"📌 <i>{sec_name}</i>",
                            parse_mode="HTML"
                        )
                    except TelegramBadRequest:
                        pass

                    img_bytes = await generate_image_gemini(keyword)
                    if img_bytes and content_paragraphs:
                        insert_idx = min((img_i + 1) * step, len(content_paragraphs) - 1)
                        p = content_paragraphs[insert_idx]

                        # Add spacing before image
                        run = p.add_run()
                        run.add_break()

                        # Add image
                        run.add_picture(_io.BytesIO(img_bytes), width=DocxInches(4.5))

                        # Add caption
                        cap_run = p.add_run()
                        cap_run.add_break()
                        images_done += 1
                        cap_run.text = f"Rasm {images_done}. {sec_name}"
                        cap_run.font.size = DocxPt(10)
                        cap_run.font.italic = True
                        cap_run.font.color.rgb = DocxRGBColor(100, 100, 100)

                    await asyncio.sleep(0.3)

                buf = _io.BytesIO()
                doc.save(buf)
                docx_bytes = buf.getvalue()

                if images_done > 0:
                    try:
                        await wait_msg.edit_text(f"✅ <b>{images_done} ta rasm qo'shildi!</b>\n📄 DOCX yakunlanmoqda...", parse_mode="HTML")
                    except TelegramBadRequest:
                        pass

            except Exception as img_err:
                logger.error(f"AI image insert error: {img_err}")

        if free_trial:
            await mark_free_used(db_user.id)
        elif db_user.id not in ADMIN_IDS:
            await deduct_balance(db_user.id, price)
        
        real_service = data.get("service_type", "essay")
        await create_request(user_id=db_user.id, service_type=real_service, topic=topic, options={"pages": pages, "ai_images": ai_images_count})
        
        await wait_msg.delete()
        prefix = "Mustaqil_ish_" if real_service == "mustaqil" else "Referat_"
        file_name = f"{prefix}{topic[:20]}.docx"
        await message.answer_document(
            document=BufferedInputFile(docx_bytes, filename=file_name),
            caption=t("doc_done", lang, topic=topic, pages=pages, price=format_price(price)),
            parse_mode="HTML",
            reply_markup=main_menu_kb(lang)
        )
        # Ask for rating
        from handlers.start import rating_kb
        service = data.get("service_type", "referat")
        rate_text = {"uz": "Iltimos, ishimizni baholang:", "ru": "Пожалуйста, оцените нашу работу:", "en": "Please rate our work:"}
        await message.answer(rate_text.get(lang, rate_text["uz"]), reply_markup=rating_kb(service))
        await state.clear()
    except Exception as e:
        logger.error(f"Document generation error: {e}")
        await message.answer(t("pres_error", lang, error=str(e)), parse_mode="HTML")
        await state.clear()

@router.callback_query(F.data == "doc_cancel_gen")
async def cancel_doc_gen(callback: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    cancel_flag = data.get("cancel_flag")
    if cancel_flag and isinstance(cancel_flag, dict):
        cancel_flag["cancelled"] = True
    await state.clear()
    await callback.message.edit_text(t("doc_cancel_gen", "uz"))
    await callback.answer()

@router.message(F.text == "❌ Bekor qilish", DocumentStates.reviewing_summary)
async def cancel_referat(message: Message, state: FSMContext):
    await state.clear()
    await message.answer(t("cancelled", "uz"), reply_markup=main_menu_kb())

