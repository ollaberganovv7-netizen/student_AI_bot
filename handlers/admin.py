from __future__ import annotations
import html as _html
from datetime import datetime

from aiogram import Router, F, Bot
from aiogram.filters import Command
from aiogram.types import CallbackQuery, Message
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

from database.db import (
    get_pending_payments, get_payment, update_payment_status,
    add_balance, get_user_count, get_premium_count,
    get_request_count, get_payment_count, get_recent_requests, get_all_users,
    set_balance
)
from keyboards.admin_kb import admin_main_kb, payment_action_kb, admin_back_kb
from keyboards.main_kb import main_menu_kb
from config import ADMIN_IDS, PRICING

router = Router()


def is_admin(user_id: int) -> bool:
    return user_id in ADMIN_IDS


class AdminStates(StatesGroup):
    broadcasting = State()
    adding_balance_id = State()
    adding_balance_amount = State()
    searching_user_id = State()
    # Kurs ishi FSM
    kurs_mavzu = State()
    kurs_fan = State()
    kurs_talaba = State()
    kurs_guruh = State()
    kurs_rahbar = State()
    # AI Rasm FSM
    ai_rasm_waiting_file = State()


# ─── Admin guard filter ───────────────────────────────────────────────────────

@router.message(Command("admin"))
async def cmd_admin(message: Message):
    if not is_admin(message.from_user.id):
        await message.answer("⛔ Sizda admin huquqlari yo'q.")
        return
    await message.answer(
        "🛡️ <b>Admin panel</b>\n\nXizmatni tanlang:",
        reply_markup=admin_main_kb(),
        parse_mode="HTML",
    )

@router.message(Command("zero"))
async def admin_zero_balance_cmd(message: Message):
    if not is_admin(message.from_user.id):
        return
    args = message.text.split()
    if len(args) < 2:
        await message.answer("⚠️ Format: /zero <user_id>")
        return
    
    try:
        user_id = int(args[1])
        user = await set_balance(user_id, 0)
        if user:
            await message.answer(f"✅ Foydalanuvchi {user_id} balansi nolga tenglashtirildi.")
            try:
                await message.bot.send_message(user_id, "⚠️ Sizning balansingiz admin tomonidan nolga tenglashtirildi.")
            except:
                pass
        else:
            await message.answer("❌ Foydalanuvchi topilmadi.")
    except ValueError:
        await message.answer("❌ Noto'g'ri ID formati.")

@router.callback_query(F.data == "admin:back")
async def admin_back(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return
    text = "🛡️ <b>Admin panel</b>\n\nXizmatni tanlang:"
    markup = admin_main_kb()

    try:
        if callback.message.photo or callback.message.document:
            await callback.message.delete()
            await callback.message.answer(text, reply_markup=markup, parse_mode="HTML")
        else:
            await callback.message.edit_text(text, reply_markup=markup, parse_mode="HTML")
    except Exception as e:
        # Fallback if edit fails
        await callback.message.delete()
        await callback.message.answer(text, reply_markup=markup, parse_mode="HTML")
        
    await callback.answer()


# ─── Pending Payments ─────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin:pending_payments")
async def admin_pending_payments(callback: CallbackQuery, bot: Bot):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    payments = await get_pending_payments()

    if not payments:
        await callback.message.edit_text(
            "✅ <b>Kutayotgan to'lovlar yo'q.</b>",
            reply_markup=admin_back_kb(),
            parse_mode="HTML",
        )
        await callback.answer()
        return

    await callback.message.edit_text(
        f"💳 <b>{len(payments)} ta kutayotgan to'lov:</b>",
        reply_markup=admin_back_kb(),
        parse_mode="HTML",
    )

    for payment in payments:
        try:
            await bot.send_photo(
                chat_id=callback.from_user.id,
                photo=payment.screenshot_file_id,
                caption=(
                    f"🔑 To'lov ID: #{payment.id}\n"
                    f"👤 Foydalanuvchi ID: <code>{payment.user_id}</code>\n"
                    f"📦 Paket: {payment.package}\n"
                    f"💰 Summa: {payment.amount:,} so'm\n"
                    f"📅 Sana: {payment.created_at.strftime('%d.%m.%Y %H:%M')}"
                ),
                reply_markup=payment_action_kb(payment.id, payment.user_id),
                parse_mode="HTML",
            )
        except Exception:
            pass

    await callback.answer()


# ─── Approve payment ──────────────────────────────────────────────────────────

@router.callback_query(F.data.startswith("admin:approve:"))
async def admin_approve_payment(callback: CallbackQuery, bot: Bot):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    parts = callback.data.split(":")
    payment_id = int(parts[2])
    user_id = int(parts[3])

    payment = await get_payment(payment_id)
    if not payment:
        await callback.answer("To'lov topilmadi", show_alert=True)
        return

    if payment.status != "pending":
        await callback.answer(f"Bu to'lov allaqachon: {payment.status}", show_alert=True)
        return

    await update_payment_status(payment_id, "approved", "Admin tomonidan tasdiqlandi")
    await add_balance(user_id, payment.amount)

    await callback.message.edit_caption(
        caption=callback.message.caption + "\n\n✅ <b>TASDIQLANDI</b>",
        parse_mode="HTML",
    )
    await callback.answer("✅ To'lov tasdiqlandi!")

    # Notify user — send them back to create with Yaratish button
    try:
        from aiogram.types import ReplyKeyboardMarkup, KeyboardButton
        create_kb = ReplyKeyboardMarkup(
            keyboard=[
                [KeyboardButton(text="✅ Yaratish")],
                [KeyboardButton(text="❌ Bekor qilish")]
            ],
            resize_keyboard=True
        )
        await bot.send_message(
            chat_id=user_id,
            text=(
                f"🎉 <b>Tabriklaymiz! To'lovingiz tasdiqlandi.</b>\n\n"

                f"💰 Balansingizga <b>{payment.amount:,} so'm</b> qo'shildi.\n\n"

                f"✅ <b>Balansingiz yetarli!</b>\n"
                f"<b>✅ Yaratish</b> tugmasini bosing!"
            ),
            reply_markup=create_kb,
            parse_mode="HTML",
        )
    except Exception:
        pass


# ─── Reject payment ───────────────────────────────────────────────────────────

@router.callback_query(F.data.startswith("admin:reject:"))
async def admin_reject_payment(callback: CallbackQuery, bot: Bot):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    parts = callback.data.split(":")
    payment_id = int(parts[2])
    user_id = int(parts[3])

    payment = await get_payment(payment_id)
    if not payment:
        await callback.answer("To'lov topilmadi", show_alert=True)
        return

    if payment.status != "pending":
        await callback.answer(f"Bu to'lov allaqachon: {payment.status}", show_alert=True)
        return

    await update_payment_status(payment_id, "rejected", "Admin tomonidan rad etildi")

    await callback.message.edit_caption(
        caption=callback.message.caption + "\n\n❌ <b>RAD ETILDI</b>",
        parse_mode="HTML",
    )
    await callback.answer("❌ To'lov rad etildi!")

    try:
        await bot.send_message(
            chat_id=user_id,
            text=(
                "❌ <b>Afsuski, to'lovingiz tasdiqlanmadi.</b>\n\n"
                "Sabab: Chek aniq emas yoki to'lov summasi noto'g'ri.\n"
                "Qayta to'lov qilib, chek yuboring."
            ),
            reply_markup=main_menu_kb(),
            parse_mode="HTML",
        )
    except Exception:
        pass


# ─── Statistics ───────────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin:stats")
async def admin_stats(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    total_users = await get_user_count()
    total_requests = await get_request_count()
    approved_payments = await get_payment_count()

    text = (
        f"📊 <b>Bot statistikasi</b>\n\n"
        f"👥 Jami foydalanuvchilar: <b>{total_users}</b>\n"
        f"📄 Jami so'rovlar: <b>{total_requests}</b>\n"
        f"✅ Tasdiqlangan to'lovlar: <b>{approved_payments}</b>\n"
    )

    await callback.message.edit_text(text, reply_markup=admin_back_kb(), parse_mode="HTML")
    await callback.answer()


# ─── Pricing list ─────────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin:pricing")
async def admin_pricing(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    lines = [f"💰 <b>Joriy narxlar:</b>\n"]
    for key, price in PRICING.items():
        lines.append(f"• {key}: <b>{price:,} so'm</b>")

    await callback.message.edit_text(
        "\n".join(lines),
        reply_markup=admin_back_kb(),
        parse_mode="HTML",
    )
    await callback.answer()


# ─── Recent requests ──────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin:recent_requests")
async def admin_recent_requests(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    requests = await get_recent_requests(limit=10)
    if not requests:
        await callback.message.edit_text(
            "📋 So'rovlar yo'q.",
            reply_markup=admin_back_kb(),
            parse_mode="HTML",
        )
        await callback.answer()
        return

    lines = ["📋 <b>So'nggi 10 ta so'rov:</b>\n"]
    for req in requests:
        lines.append(
            f"• [{req.created_at.strftime('%d.%m %H:%M')}] "
            f"<code>{req.user_id}</code> → {req.service_type}: <i>{req.topic[:40]}</i>"
        )

    await callback.message.edit_text(
        "\n".join(lines),
        reply_markup=admin_back_kb(),
        parse_mode="HTML",
    )
    await callback.answer()


# ─── Reactions Stats ──────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin:reactions")
async def admin_reactions(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    from sqlalchemy import select, func, and_, desc
    from database.db import async_session
    from database.models import PostReaction

    async with async_session() as session:
        # Get recent reactions grouped by message
        result = await session.execute(
            select(
                PostReaction.chat_id,
                PostReaction.message_id,
                PostReaction.emoji,
                PostReaction.user_name,
                PostReaction.created_at,
            )
            .order_by(desc(PostReaction.created_at))
            .limit(50)
        )
        reactions = result.all()

    if not reactions:
        await callback.message.edit_text(
            "📭 <b>Hech qanday reaktsiya yo'q.</b>\n\n"
            "Kanal postlariga hech kim reaktsiya qo'ymagan.",
            reply_markup=admin_back_kb(),
            parse_mode="HTML",
        )
        await callback.answer()
        return

    # Group by message
    by_msg: dict[int, dict] = {}
    for chat_id, msg_id, emoji, user_name, created_at in reactions:
        if msg_id not in by_msg:
            by_msg[msg_id] = {"chat_id": chat_id, "reactions": []}
        by_msg[msg_id]["reactions"].append(f"{emoji} {user_name or 'User'}")

    lines = ["🔥 <b>So'nggi reaktsiyalar:</b>\n"]
    for msg_id, data in list(by_msg.items())[:10]:
        lines.append(f"📌 <b>Post #{msg_id}:</b>")
        for r in data["reactions"][:10]:
            lines.append(f"  • {r}")
        if len(data["reactions"]) > 10:
            lines.append(f"  ... +{len(data['reactions'])-10} ta")
        lines.append("")

    text = "\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "..."

    await callback.message.edit_text(
        text,
        reply_markup=admin_back_kb(),
        parse_mode="HTML",
    )
    await callback.answer()


# ─── Add Balance Manually ───────────────────────────────────────────────────

@router.callback_query(F.data == "admin:add_balance")
async def admin_add_balance_start(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return
    await state.set_state(AdminStates.adding_balance_id)
    await callback.message.edit_text(
        "➕ <b>Balans qo'shish</b>\n\n"
        "👤 Foydalanuvchining ID raqamini yozing:\n"
        "<i>(Bekor qilish uchun /cancel)</i>",
        reply_markup=admin_back_kb(),
        parse_mode="HTML",
    )
    await callback.answer()

@router.message(AdminStates.adding_balance_id)
async def admin_add_balance_get_id(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    if message.text == "/cancel":
        await state.clear()
        await message.answer("❌ Bekor qilindi.", reply_markup=admin_back_kb())
        return

    try:
        user_id = int(message.text.strip())
    except ValueError:
        await message.answer("⚠️ Iltimos, faqat ID raqam yozing.")
        return

    await state.update_data(target_user_id=user_id)
    await state.set_state(AdminStates.adding_balance_amount)
    await message.answer(
        f"✅ ID: <code>{user_id}</code> qabul qilindi.\n\n"
        "💰 Qancha summa qo'shmoqchisiz?\n"
        "<i>(Masalan: 5000 yoki 10000)</i>\n"
        "<i>(Bekor qilish uchun /cancel)</i>",
        parse_mode="HTML",
        reply_markup=admin_back_kb()
    )

@router.message(AdminStates.adding_balance_amount)
async def admin_add_balance_finish(message: Message, state: FSMContext, bot: Bot):
    if not is_admin(message.from_user.id):
        return
    if message.text == "/cancel":
        await state.clear()
        await message.answer("❌ Bekor qilindi.", reply_markup=admin_back_kb())
        return

    try:
        amount = int(message.text.strip().replace(" ", "").replace(",", ""))
    except ValueError:
        await message.answer("⚠️ Iltimos, faqat raqam yozing.")
        return

    data = await state.get_data()
    user_id = data.get("target_user_id")

    try:
        await add_balance(user_id, amount)
        await state.clear()
        await message.answer(
            f"✅ <b>Muvaffaqiyatli!</b>\n\n"
            f"👤 ID: <code>{user_id}</code>\n"
            f"💰 Qo'shildi: {amount:,} so'm",
            parse_mode="HTML",
            reply_markup=admin_back_kb()
        )
        
        # Notify user
        try:
            await bot.send_message(
                chat_id=user_id,
                text=(
                    f"🎉 <b>Tabriklaymiz! Admin tomonidan balansingiz to'ldirildi.</b>\n\n"

                    f"💰 Balansingizga <b>{amount:,} so'm</b> qo'shildi.\n"
                    f"Xizmatlardan foydalanishingiz mumkin!"
                ),
                reply_markup=main_menu_kb(),
                parse_mode="HTML"
            )
        except Exception:
            await message.answer("⚠️ Foydalanuvchiga xabar yuborib bo'lmadi (botni bloklagan bo'lishi mumkin).")

    except Exception as e:
        await message.answer(f"❌ Xatolik yuz berdi: {e}")

# ─── Broadcast ────────────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin:broadcast")
async def admin_broadcast_start(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    await state.set_state(AdminStates.broadcasting)
    await callback.message.edit_text(
        "📢 <b>Xabar yuborish</b>\n\n"
        "Barcha foydalanuvchilarga yuboriladigan xabarni yozing:\n"
        "<i>(Bekor qilish uchun /cancel)</i>",
        reply_markup=admin_back_kb(),
        parse_mode="HTML",
    )
    await callback.answer()


@router.message(AdminStates.broadcasting)
async def admin_broadcast_send(message: Message, state: FSMContext, bot: Bot):
    if not is_admin(message.from_user.id):
        return

    if message.text and message.text.strip() == "/cancel":
        await state.clear()
        await message.answer("❌ Xabar bekor qilindi.", reply_markup=admin_back_kb())
        return

    await state.clear()
    users = await get_all_users()
    sent = 0
    failed = 0

    status_msg = await message.answer(f"📤 Yuborilmoqda... 0/{len(users)}")

    broadcast_text = _html.escape(message.text or message.caption or '')

    for i, user in enumerate(users):
        try:
            if message.photo:
                await bot.send_photo(
                    chat_id=user.id,
                    photo=message.photo[-1].file_id,
                    caption=f"📢 <b>Bot xabari:</b>\n\n{broadcast_text}",
                    parse_mode="HTML",
                )
            elif message.video:
                await bot.send_video(
                    chat_id=user.id,
                    video=message.video.file_id,
                    caption=f"📢 <b>Bot xabari:</b>\n\n{broadcast_text}",
                    parse_mode="HTML",
                )
            elif message.document:
                await bot.send_document(
                    chat_id=user.id,
                    document=message.document.file_id,
                    caption=f"📢 <b>Bot xabari:</b>\n\n{broadcast_text}",
                    parse_mode="HTML",
                )
            else:
                await bot.send_message(
                    chat_id=user.id,
                    text=f"📢 <b>Bot xabari:</b>\n\n{broadcast_text}",
                    parse_mode="HTML",
                )
            sent += 1
        except Exception:
            failed += 1

        if (i + 1) % 20 == 0:
            try:
                await status_msg.edit_text(f"📤 Yuborilmoqda... {i+1}/{len(users)}")
            except Exception:
                pass

    await status_msg.edit_text(
        f"✅ <b>Xabar yuborildi!</b>\n\n"
        f"✅ Muvaffaqiyatli: {sent}\n"
        f"❌ Xatolik: {failed}",
        parse_mode="HTML",
    )

# ─── Search User by ID ───────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin:search_user")
async def admin_search_user_start(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    await state.set_state(AdminStates.searching_user_id)
    await callback.message.edit_text(
        "🔍 <b>Foydalanuvchi qidirish</b>\n\n"
        "Foydalanuvchi Telegram ID sini kiriting:\n"
        "<i>(Bekor qilish uchun /cancel)</i>",
        reply_markup=admin_back_kb(),
        parse_mode="HTML",
    )
    await callback.answer()


@router.message(AdminStates.searching_user_id)
async def admin_search_user_result(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return

    if message.text and message.text.strip() == "/cancel":
        await state.clear()
        await message.answer("❌ Qidirish bekor qilindi.", reply_markup=admin_back_kb())
        return

    try:
        user_id = int(message.text.strip())
    except ValueError:
        await message.answer("⚠️ Noto'g'ri ID! Faqat raqam kiriting.")
        return

    from database.db import get_user
    from utils.helpers import format_price

    user = await get_user(user_id)
    await state.clear()

    if not user:
        await message.answer(
            f"❌ <b>Foydalanuvchi topilmadi!</b>\n\n"
            f"ID: {user_id}\n"
            f"Bunday ID bilan foydalanuvchi yo'q.",
            reply_markup=admin_back_kb(),
            parse_mode="HTML",
        )
        return

    premium_status = "✅ Faol" if user.is_premium and user.premium_expires and user.premium_expires > datetime.utcnow() else "❌ Yo'q"
    premium_expiry = user.premium_expires.strftime("%Y-%m-%d") if user.premium_expires else "—"
    none_label = "Yo'q"

    safe_name = _html.escape(user.full_name or none_label)
    safe_username = _html.escape(user.username or none_label)
    await message.answer(
        f"👤 <b>Foydalanuvchi ma'lumotlari</b>\n\n"
        f"🆔 <b>ID:</b> {user.id}\n"
        f"👤 <b>Ism:</b> {safe_name}\n"
        f"📛 <b>Username:</b> @{safe_username}\n"
        f"💳 <b>Balans:</b> {format_price(user.balance or 0)}\n"
        f"💎 <b>Premium:</b> {premium_status}\n"
        f"📅 <b>Premium muddati:</b> {premium_expiry}\n"
        f"🆓 <b>Bepul ishlatgan:</b> {'Ha' if user.free_used else none_label}\n"
        f"🕐 <b>Oxirgi faollik:</b> {user.last_active.strftime('%Y-%m-%d %H:%M') if user.last_active else none_label}",
        reply_markup=admin_back_kb(),
        parse_mode="HTML",
    )

# ─── Lab Generator ───────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin:lab_gen")
async def admin_lab_gen(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    from keyboards.admin_kb import lab_variant_kb
    await callback.message.edit_text(
        "🔬 <b>Lab Generator</b>\n\n"
        "Praktik ish №2 — GPIO, knopka bilan LED boshqarish\n\n"
        "Variant raqamini tanlang (1-30):\n"
        "<i>Rezistor va chastota avtomatik o'zgaradi</i>",
        reply_markup=lab_variant_kb(),
        parse_mode="HTML",
    )
    await callback.answer()


@router.callback_query(F.data.startswith("admin:lab_v:"))
async def admin_lab_variant(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    variant = int(callback.data.split(":")[2])

    from services.lab_generator import generate_lab2, VARIANTS
    resistor, freq = VARIANTS[variant]

    await callback.message.edit_text(
        f"⏳ Variant {variant} tayyorlanmoqda...\n"
        f"Rezistor: {resistor} Ω | Chastota: {freq} MHz",
        parse_mode="HTML",
    )

    try:
        buf = generate_lab2(variant)
        from aiogram.types import BufferedInputFile
        doc_file = BufferedInputFile(
            buf.read(),
            filename=f"Lab2_variant_{variant}.docx",
        )
        await callback.message.answer_document(
            doc_file,
            caption=(
                f"🔬 <b>Praktik ish №2 — Variant {variant}</b>\n\n"
                f"📌 Rezistor: <b>{resistor} Ω</b>\n"
                f"📌 Chastota: <b>{freq} MHz</b>\n\n"
                f"Fayl tayyor! Rasmlar va jadval saqlangan."
            ),
            parse_mode="HTML",
        )
        from keyboards.admin_kb import lab_variant_kb
        await callback.message.edit_text(
            f"✅ Variant {variant} yuborildi!\n\n"
            "Boshqa variant tanlang yoki orqaga qayting:",
            reply_markup=lab_variant_kb(),
            parse_mode="HTML",
        )
    except Exception as e:
        await callback.message.edit_text(
            f"❌ Xatolik: {e}",
            reply_markup=admin_back_kb(),
            parse_mode="HTML",
        )

# ─── Kurs Ishi Generator ────────────────────────────────────────────────────

@router.callback_query(F.data == "admin:kurs_ishi")
async def admin_kurs_ishi_start(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    await state.set_state(AdminStates.kurs_mavzu)
    await callback.message.edit_text(
        "📝 <b>Kurs ishi generator</b>\n\n"
        "1/5 — Kurs ishi <b>mavzusini</b> yozing:",
        parse_mode="HTML",
    )
    await callback.answer()


@router.message(AdminStates.kurs_mavzu)
async def kurs_get_mavzu(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.update_data(mavzu=message.text.strip())
    await state.set_state(AdminStates.kurs_fan)
    await message.answer(
        "📝 2/5 — <b>Fan nomini</b> yozing:\n"
        "<i>Masalan: Informatika, Iqtisodiyot, Huquqshunoslik...</i>",
        parse_mode="HTML",
    )

@router.message(AdminStates.kurs_fan)
async def kurs_get_fan(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.update_data(fan=message.text.strip())
    await state.set_state(AdminStates.kurs_talaba)
    await message.answer(
        "📝 3/5 — <b>Talaba F.I.Sh.</b> ni yozing:",
        parse_mode="HTML",
    )

@router.message(AdminStates.kurs_talaba)
async def kurs_get_talaba(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.update_data(talaba=message.text.strip())
    await state.set_state(AdminStates.kurs_guruh)
    await message.answer(
        "📝 4/5 — <b>Guruh raqamini</b> yozing:\n"
        "<i>Masalan: 110-22, AI-23...</i>",
        parse_mode="HTML",
    )

@router.message(AdminStates.kurs_guruh)
async def kurs_get_guruh(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.update_data(guruh=message.text.strip())
    await state.set_state(AdminStates.kurs_rahbar)
    await message.answer(
        "📝 5/5 — <b>Ilmiy rahbar F.I.Sh.</b> ni yozing:",
        parse_mode="HTML",
    )

@router.message(AdminStates.kurs_rahbar)
async def kurs_get_rahbar(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return

    data = await state.get_data()
    await state.clear()

    mavzu = data["mavzu"]
    fan = data["fan"]
    talaba = data["talaba"]
    guruh = data["guruh"]
    rahbar = message.text.strip()

    await message.answer("⏳ Kurs ishi tayyorlanmoqda...")

    try:
        from services.kurs_ishi_generator import generate_kurs_ishi
        buf = generate_kurs_ishi(
            mavzu=mavzu,
            fan=fan,
            talaba=talaba,
            guruh=guruh,
            rahbar=rahbar,
        )
        from aiogram.types import BufferedInputFile
        doc_file = BufferedInputFile(
            buf.read(),
            filename=f"Kurs_ishi_{mavzu[:30].replace(' ', '_')}.docx",
        )
        await message.answer_document(
            doc_file,
            caption=(
                f"📝 <b>Kurs ishi tayyor!</b>\n\n"
                f"📌 Mavzu: <b>{mavzu}</b>\n"
                f"📌 Fan: <b>{fan}</b>\n"
                f"📌 Talaba: <b>{talaba}</b>\n"
                f"📌 Guruh: <b>{guruh}</b>\n"
                f"📌 Rahbar: <b>{rahbar}</b>\n\n"
                f"⚠️ Asosiy qism va xulosa — shablonli. "
                f"O'z matnini yozing!"
            ),
            parse_mode="HTML",
        )
        await message.answer(
            "Yana kurs ishi yaratish yoki admin panelga qaytish:",
            reply_markup=admin_main_kb(),
        )
    except Exception as e:
        await message.answer(
            f"❌ Xatolik: {e}",
            reply_markup=admin_main_kb(),
        )

# ─── AI Rasm (faylga rasm qo'shish) ─────────────────────────────────────────

@router.callback_query(F.data == "admin:ai_rasm")
async def admin_ai_rasm_start(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        await callback.answer("⛔ Ruxsat yo'q", show_alert=True)
        return

    await state.set_state(AdminStates.ai_rasm_waiting_file)
    await callback.message.edit_text(
        "🖼 <b>AI Rasm — Faylga rasm qo'shish</b>\n\n"
        "📎 <b>Fayl yuboring</b> (.docx, .pdf yoki .txt)\n\n"
        "<i>AI faylni o'qib, har bir bo'lim uchun mavzuga mos\n"
        "rasm generatsiya qiladi va DOCX ichiga joylaydi.</i>\n\n"
        "<i>(Bekor qilish uchun /cancel)</i>",
        reply_markup=admin_back_kb(),
        parse_mode="HTML",
    )
    await callback.answer()


@router.message(AdminStates.ai_rasm_waiting_file, F.text == "/cancel")
async def admin_ai_rasm_cancel(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.clear()
    await message.answer("❌ Bekor qilindi.", reply_markup=admin_back_kb())


@router.message(AdminStates.ai_rasm_waiting_file, F.document)
async def admin_ai_rasm_receive_file(message: Message, state: FSMContext, bot: Bot):
    if not is_admin(message.from_user.id):
        return

    import io
    import json
    import asyncio

    doc = message.document
    file_name = doc.file_name or "document"
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

    if ext not in ("docx", "pdf", "txt", "doc", "md"):
        await message.answer(
            "⚠️ Faqat <b>.docx, .pdf, .txt</b> formatlarni qo'llab-quvvatlaydi.",
            parse_mode="HTML",
        )
        return

    status_msg = await message.answer(
        "⏳ <b>Fayl o'qilmoqda...</b>",
        parse_mode="HTML",
    )

    try:
        # ── Step 1: Download and read the file ──
        file_obj = await bot.get_file(doc.file_id)
        buf = io.BytesIO()
        await bot.download_file(file_obj.file_path, buf)
        buf.seek(0)
        file_bytes = buf.read()

        from services.file_reader_service import read_document
        text = read_document(file_bytes, file_name)

        if not text or len(text.strip()) < 50:
            await status_msg.edit_text(
                "❌ Fayldan matn chiqarib bo'lmadi yoki fayl bo'sh.",
            )
            await state.clear()
            return

        text_truncated = text[:12000]
        word_count = len(text.split())

        await status_msg.edit_text(
            f"✅ <b>Fayl o'qildi!</b>\n"
            f"📝 Hajm: ~{word_count} so'z\n\n"
            f"🧠 <b>AI bo'limlarni tahlil qilmoqda...</b>",
            parse_mode="HTML",
        )

        # ── Step 2: AI analyzes sections and generates image keywords ──
        from services.ai_service import _call_ai

        analysis_prompt = (
            f"Quyidagi akademik hujjat matnini o'qi va tahlil qil.\n\n"
            f"MATN:\n{text_truncated}\n\n"
            f"VAZIFA: Har bir asosiy bo'lim/mavzu uchun 1 ta rasm kerak.\n"
            f"Har bir rasm uchun INGLIZ tilida aniq, tasviriy kalit so'z yoz.\n"
            f"Kalit so'z photo-realistic rasm generatsiya qilish uchun ishlatiladi.\n\n"
            f"FAQAT JSON formatda javob ber (boshqa hech narsa qo'shma):\n"
            f'[\n'
            f'  {{"section": "Bo\'lim nomi", "image_keyword": "professional photo of economics graphs and charts"}},\n'
            f'  {{"section": "Ikkinchi bo\'lim", "image_keyword": "modern university classroom with students"}},\n'
            f'  ...\n'
            f']\n\n'
            f"QOIDALAR:\n"
            f"- Kamida 3 ta, ko'pi bilan 10 ta rasm bo'lsin\n"
            f"- image_keyword 5-15 so'zdan iborat, aniq va tasviriy bo'lsin\n"
            f"- Rasmlar mavzuga TO'LIQMOS bo'lishi SHART\n"
            f"- Faqat JSON array qaytar"
        )

        raw_analysis = await _call_ai(
            [{"role": "user", "content": analysis_prompt}],
            max_tokens=2000,
            temperature=0.5,
            json_mode=True,
        )

        import re
        clean = raw_analysis.replace("```json", "").replace("```", "").strip()
        sections = json.loads(clean)

        if not isinstance(sections, list) or len(sections) == 0:
            await status_msg.edit_text("❌ AI bo'limlarni aniqlay olmadi.")
            await state.clear()
            return

        # Limit to 10 images
        sections = sections[:10]
        total = len(sections)

        await status_msg.edit_text(
            f"✅ <b>{total} ta bo'lim aniqlandi!</b>\n\n"
            + "\n".join(f"  {i+1}. {s.get('section', '?')}" for i, s in enumerate(sections))
            + f"\n\n🎨 <b>Gemini AI rasmlar generatsiya qilmoqda...</b>\n"
            f"⏳ Har bir rasm ~10-20 soniya...",
            parse_mode="HTML",
        )

        # ── Step 3: Generate images with Gemini ──
        from services.image_search_service import generate_image_gemini

        generated_images = []  # list of (section_name, image_bytes)

        for idx, sec in enumerate(sections):
            keyword = sec.get("image_keyword", "")
            section_name = sec.get("section", f"Bo'lim {idx+1}")

            if not keyword:
                continue

            bar_done = "🟩" * (idx + 1)
            bar_left = "⬜" * (total - idx - 1)
            try:
                await status_msg.edit_text(
                    f"🎨 <b>Rasm {idx+1}/{total} generatsiya qilinmoqda</b>\n\n"
                    f"{bar_done}{bar_left}\n\n"
                    f"📌 <b>{section_name}</b>\n"
                    f"🔑 <i>{keyword}</i>",
                    parse_mode="HTML",
                )
            except:
                pass

            img_bytes = await generate_image_gemini(keyword)
            if img_bytes:
                generated_images.append((section_name, img_bytes))

            await asyncio.sleep(0.5)

        if not generated_images:
            await status_msg.edit_text(
                "❌ Hech qanday rasm generatsiya qilib bo'lmadi.\n"
                "Gemini API bilan muammo bo'lishi mumkin.",
            )
            await state.clear()
            return

        # ── Step 4: Create DOCX with images inserted ──
        try:
            await status_msg.edit_text(
                f"✅ <b>{len(generated_images)} ta rasm tayyor!</b>\n\n"
                f"📄 <b>DOCX fayl yaratilmoqda...</b>",
                parse_mode="HTML",
            )
        except:
            pass

        # If original file is DOCX — insert images into it
        # Otherwise — create a new DOCX with text + images
        from docx import Document as DocxDocument
        from docx.shared import Inches as DocxInches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if ext == "docx":
            # Insert images into original DOCX
            doc = DocxDocument(io.BytesIO(file_bytes))
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]

            if len(paragraphs) == 0:
                paragraphs = list(doc.paragraphs)

            step = max(1, len(paragraphs) // (len(generated_images) + 1))

            for img_i, (sec_name, img_bytes_data) in enumerate(generated_images):
                insert_idx = min((img_i + 1) * step, len(paragraphs) - 1)
                p = paragraphs[insert_idx]

                # Add spacing
                run = p.add_run()
                run.add_break()

                # Add image
                run.add_picture(io.BytesIO(img_bytes_data), width=DocxInches(5.0))

                # Add caption below
                caption_run = p.add_run()
                caption_run.add_break()
                caption_run.text = f"Rasm {img_i + 1}. {sec_name}"
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            # Create new DOCX for non-docx files
            doc = DocxDocument()

            # Title
            title_para = doc.add_heading(f"AI generatsiya qilingan rasmlar", level=1)
            doc.add_paragraph(f"Manba: {file_name}")
            doc.add_paragraph("")

            for img_i, (sec_name, img_bytes_data) in enumerate(generated_images):
                # Section heading
                doc.add_heading(sec_name, level=2)

                # Image
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(io.BytesIO(img_bytes_data), width=DocxInches(5.0))

                # Caption
                caption = doc.add_paragraph(f"Rasm {img_i + 1}. {sec_name}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = caption.runs[0] if caption.runs else None
                if cap_run:
                    cap_run.font.size = Pt(10)
                    cap_run.font.italic = True
                    cap_run.font.color.rgb = RGBColor(100, 100, 100)

                doc.add_paragraph("")  # spacing

        # Save DOCX
        output_buf = io.BytesIO()
        doc.save(output_buf)
        docx_bytes = output_buf.getvalue()

        # ── Step 5: Send result ──
        from aiogram.types import BufferedInputFile

        output_name = file_name.rsplit(".", 1)[0] + "_AI_rasmli.docx"
        doc_file = BufferedInputFile(docx_bytes, filename=output_name)

        await status_msg.delete()
        await message.answer_document(
            doc_file,
            caption=(
                f"🖼 <b>AI Rasm — Tayyor!</b>\n\n"
                f"📎 Manba: <i>{_html.escape(file_name)}</i>\n"
                f"🎨 Rasmlar soni: <b>{len(generated_images)}</b>\n"
                f"📄 DOCX ichiga joylandi\n\n"
                f"<i>Gemini AI tomonidan generatsiya qilindi</i>"
            ),
            parse_mode="HTML",
        )
        await message.answer(
            "Yana rasm yaratish yoki admin panelga qaytish:",
            reply_markup=admin_main_kb(),
        )
        await state.clear()

    except json.JSONDecodeError:
        await status_msg.edit_text(
            "❌ AI javobini o'qib bo'lmadi. Qayta urinib ko'ring.",
            reply_markup=admin_back_kb(),
        )
        await state.clear()
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"AI Rasm error: {e}", exc_info=True)
        try:
            await status_msg.edit_text(
                f"❌ Xatolik yuz berdi: {e}",
                reply_markup=admin_back_kb(),
            )
        except:
            await message.answer(f"❌ Xatolik: {e}", reply_markup=admin_back_kb())
        await state.clear()


@router.message(AdminStates.ai_rasm_waiting_file)
async def admin_ai_rasm_wrong_input(message: Message):
    if not is_admin(message.from_user.id):
        return
    await message.answer(
        "⚠️ Iltimos, <b>fayl</b> yuboring (.docx, .pdf, .txt)\n"
        "Yoki /cancel bilan bekor qiling.",
        parse_mode="HTML",
    )

# ─── Zero Balance Command ─────────────────────────────────────────────────────

@router.message(Command(commands=["zero", "nol"]))
async def zero_balance_command(message: Message):
    if not is_admin(message.from_user.id):
        return
        
    parts = message.text.split()
    if len(parts) != 2 or not parts[1].isdigit():
        await message.answer("⚠️ Noto'g'ri format. Foydalanish: <code>/zero ID</code>", parse_mode="HTML")
        return
        
    target_id = int(parts[1])
    from database.db import set_balance
    user = await set_balance(target_id, 0)
    
    if user:
        await message.answer(f"✅ Foydalanuvchi {target_id} ({user.full_name}) balansi 0 qilib belgilandi.")
    else:
        await message.answer("⚠️ Bunday ID ga ega foydalanuvchi topilmadi.")


